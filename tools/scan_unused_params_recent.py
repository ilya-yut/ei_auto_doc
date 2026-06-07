#!/usr/bin/env python3
"""Scan recent generated Explanation docs for params in table but not in ABAP interface."""
from __future__ import annotations

import re
from collections import defaultdict
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output"
INPUT_DIRS = [ROOT / "input", ROOT / "input" / "old"]
CUTOFF = datetime.now() - timedelta(days=14)

RE_FM = re.compile(r"FUNCTION\s+(/SKN/F_[\w.]+)", re.I)
RE_HARDCODED_SY = re.compile(r"^\s*(\w+)\s+SY-DATUM\b", re.M | re.I)


def _add_multiline_block(lines: list[str], start_idx: int, names: set[str]) -> int:
    ln = lines[start_idx]
    rest = re.sub(r"^\s*\w+_MULTY:\s*", "", ln)
    mm0 = re.match(r"(\w+)", rest.strip())
    if mm0:
        names.add(mm0.group(1).upper())
    j = start_idx + 1
    while j < len(lines):
        s = lines[j]
        if re.match(r"^\s*DATA:\s", s) or re.match(r"^\s*(DATA_|SELECT_|LV_)", s):
            break
        mm = re.match(r"^\s+(\w+)\s*[,]", s) or re.match(r"^\s+(\w+)\s+\S", s)
        if mm:
            names.add(mm.group(1).upper())
        j += 1
    return j


def parse_interface_params(text: str) -> set[str]:
    """All fields read via SELECT_SINGLE/MULTY (incl. continuation lines)."""
    names: set[str] = set()
    m = re.search(r"DATA_SINGLE:\s*([\s\S]*?)^\s*DATA_MULTY:", text, re.MULTILINE)
    if m:
        for ln in m.group(1).splitlines():
            t = ln.strip()
            if not t or t.startswith('"'):
                continue
            mm = re.match(r"^,?\s*(\w+)\s+", t)
            if mm:
                names.add(mm.group(1).upper())
    for m in re.finditer(r"^\s*DATA_SINGLE:\s*(\w+)\s+\S", text, re.MULTILINE):
        names.add(m.group(1).upper())
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        if re.match(r"^\s*DATA_MULTY:", lines[i]) or re.match(r"^\s*SELECT_MULTY:", lines[i]):
            i = _add_multiline_block(lines, i, names)
            continue
        i += 1
    i = 0
    while i < len(lines):
        ln = lines[i]
        if not re.match(r"^\s*SELECT_SINGLE:\s*", ln):
            i += 1
            continue
        tail = re.sub(r"^\s*SELECT_SINGLE:\s*", "", ln).strip()
        if tail:
            tok = tail.split(",")[0].strip()
            if tok.isidentifier():
                names.add(tok.upper())
        i += 1
        while i < len(lines):
            s = lines[i]
            st = s.strip()
            if (
                not st
                or st.startswith('"---')
                or re.match(r"^\s*DATA_SINGLE:\s*SW_DEST", s)
                or re.match(r"^\s*DATA:\s", s)
            ):
                break
            if len(s) - len(s.lstrip()) < 10:
                break
            mm = re.match(r"^\s*(\w+)\s*,", s) or re.match(r"^\s*(\w+)\s*\.\s*$", s)
            if mm:
                names.add(mm.group(1).upper())
            i += 1
    if re.search(r"\bSW_DEST\b", text, re.I):
        names.add("SW_DEST")
    return names


def build_code_map():
    fm_map: dict[str, Path] = {}
    stem_map: dict[str, Path] = {}
    for folder in INPUT_DIRS:
        if not folder.exists():
            continue
        for p in list(folder.glob("Code_*.txt")) + list(folder.glob("Code _*.txt")):
            try:
                text = p.read_text(encoding="utf-8", errors="replace")
            except OSError:
                continue
            m = RE_FM.search(text[:5000])
            if m:
                fm_map[m.group(1).upper().rstrip(".")] = p
            stem = p.stem
            for pref in ("Code_", "Code "):
                if stem.startswith(pref):
                    stem = stem[len(pref) :]
            stem_map[stem.upper()] = p
    return fm_map, stem_map


def parse_code_usage(text: str) -> dict:
    iface = parse_interface_params(text)
    hardcoded_sy: set[str] = set()
    for m in RE_HARDCODED_SY.finditer(text):
        hardcoded_sy.add(m.group(1).upper())
    return {"iface": iface, "hardcoded_sy": hardcoded_sy}


def resolve_code_path(fm: str | None, doc_stem: str, fm_map, stem_map):
    if fm and fm.upper() in fm_map:
        return fm_map[fm.upper()]
    s = doc_stem
    for pref in ("Explanation_", "Explanation "):
        if s.startswith(pref):
            s = s[len(pref) :]
    for c in (s.upper(), s.replace(" ", "_").upper()):
        if c in stem_map:
            return stem_map[c]
    m = re.search(r"(SKN_S_SW_[\d_]+[A-Z0-9_]*)", s.upper())
    if m and m.group(1) in stem_map:
        return stem_map[m.group(1)]
    return None


def parse_md(path: Path):
    text = path.read_text(encoding="utf-8", errors="replace")
    params = []
    in_sec = False
    for line in text.splitlines():
        if line.strip().lower().startswith("### parameters reference table"):
            in_sec = True
            continue
        if in_sec and line.startswith("### "):
            break
        if in_sec and line.startswith("|") and not line.startswith("|---"):
            parts = [x.strip() for x in line.strip("|").split("|")]
            if parts and parts[0].isdigit() and len(parts) >= 2:
                p = parts[1].strip().upper()
                if p:
                    params.append(p)
    fm = None
    m = RE_FM.search(text)
    if m:
        fm = m.group(1).upper().rstrip(".")
    return params, fm


def parse_docx(path: Path):
    doc = Document(str(path))
    params = []
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        hdr = [c.text.strip().lower() for c in tbl.rows[0].cells]
        pcol = None
        for name in ("parameter", "field"):
            if name in hdr:
                pcol = hdr.index(name)
                break
        if pcol is None or "structure name" in hdr:
            continue
        got = []
        for row in tbl.rows[1:]:
            if pcol < len(row.cells):
                p = row.cells[pcol].text.strip().upper()
                if p and p not in {"PARAMETER", "FIELD"}:
                    got.append(p)
        if len(got) >= 3:
            params = got
            break
    text = "\n".join(p.text for p in doc.paragraphs)
    fm = None
    m = RE_FM.search(text)
    if m:
        fm = m.group(1).upper().rstrip(".")
    return params, fm


def main() -> None:
    fm_map, stem_map = build_code_map()
    results = []
    no_code = []
    seen_stems: set[str] = set()

    files = sorted(OUTPUT.glob("Explanation_*.md"))
    for fp in files:
        if fp.name.startswith("~$") or " - Copy" in fp.name:
            continue
        mt = datetime.fromtimestamp(fp.stat().st_mtime)
        if mt < CUTOFF:
            continue
        stem_key = fp.stem
        if stem_key in seen_stems:
            continue
        seen_stems.add(stem_key)

        params, fm = parse_md(fp)
        code_path = resolve_code_path(fm, fp.stem, fm_map, stem_map)
        if not code_path:
            no_code.append(fp.name)
            continue

        usage = parse_code_usage(code_path.read_text(encoding="utf-8", errors="replace"))
        iface_set = usage["iface"]
        not_in_iface = sorted(p for p in params if p not in iface_set)
        hardcoded = sorted(p for p in params if p in usage["hardcoded_sy"])

        if not_in_iface or hardcoded:
            results.append(
                {
                    "file": fp.name,
                    "mtime": mt.strftime("%Y-%m-%d"),
                    "fm": fm,
                    "code": code_path.name,
                    "not_in_interface": not_in_iface,
                    "hardcoded_sy_datum": hardcoded,
                }
            )

    print("=== Recent generated docs with unused-like parameters ===")
    print(f"Cutoff: {CUTOFF.date()} (last 14 days)")
    print(f"Docs with issues: {len(results)}")
    print(f"No code match: {len(no_code)}")
    print()

    freq_not: dict[str, int] = defaultdict(int)
    freq_hard: dict[str, int] = defaultdict(int)
    for r in results:
        for p in r["not_in_interface"]:
            freq_not[p] += 1
        for p in r["hardcoded_sy_datum"]:
            freq_hard[p] += 1

    print("A) In Parameters table but NOT in DATA_SINGLE/MULTY/SELECT (output/structure-only):")
    for p, n in sorted(freq_not.items(), key=lambda x: (-x[1], x[0])):
        print(f"   {p}: {n} doc(s)")

    print()
    print("B) In table AND in DATA_MULTY but hardcoded to SY-DATUM (not user-configurable):")
    for p, n in sorted(freq_hard.items(), key=lambda x: (-x[1], x[0])):
        print(f"   {p}: {n} doc(s)")

    print()
    print("--- Per document ---")
    for r in sorted(results, key=lambda x: (x["mtime"], x["file"])):
        bits = []
        if r["not_in_interface"]:
            ni = r["not_in_interface"]
            s = ",".join(ni[:15])
            if len(ni) > 15:
                s += f",...+{len(ni)-15}"
            bits.append(f"not_in_iface=[{s}]")
        if r["hardcoded_sy_datum"]:
            bits.append("hardcoded_SY-DATUM=" + ",".join(r["hardcoded_sy_datum"]))
        print(f"{r['mtime']}  {r['file']}")
        print("    " + "; ".join(bits))

    if no_code:
        print()
        print("No code file (skipped):")
        for n in sorted(no_code):
            print(f"  {n}")


if __name__ == "__main__":
    main()
