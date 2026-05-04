"""
MD → DOCX conversion for the EI doc pipeline.
Uses markdown → HTML (markdown + tables + fenced_code) → BeautifulSoup → python-docx.
No Pandoc or reference document; formatting is applied in code.
Adapted from ei docs / doc generator scripts 5 / scripts / md_to_docx_converter.py.
"""

import re
from pathlib import Path

# XML 1.0 allows only #x9, #xA, #xD as control chars; python-docx fails on NULL and other control chars
_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")


def _md_docx_use_plain_list_items(current_section: str, *, options_ul: bool) -> bool:
    """
    When True, <ul><li> items are rendered as normal paragraphs (no Word list bullet).
    Used for parameter Options lists and for list-heavy Parameter subsections.
    """
    if options_ul:
        return True
    s = (current_section or "").lower()
    return any(
        frag in s
        for frag in (
            "parameter configuration guidelines",
            "parameter relationship",
            "default value",
            "parameters reference table",
            "practical example",
        )
    )


def _md_docx_no_list_styles(
    current_section: str, no_bullet_sections: list[str], *, options_ul: bool
) -> bool:
    """Use normal paragraphs instead of Word list styles for bullets/dash-lines in these sections."""
    if current_section in no_bullet_sections:
        return True
    return _md_docx_use_plain_list_items(current_section, options_ul=options_ul)


def _sanitize_xml_text(s: str) -> str:
    """Remove/replace control characters so strings are XML-compatible for python-docx."""
    if not s:
        return s
    return _CONTROL_CHARS.sub("", s)


def _parse_markdown_to_elements(md_content: str):
    """Parse markdown content into structured elements (HTML soup)."""
    import markdown
    from bs4 import BeautifulSoup

    md = markdown.Markdown(extensions=["tables", "fenced_code"])
    html = md.convert(md_content)
    return BeautifulSoup(html, "html.parser")


def _add_code_block(doc, code_text: str) -> None:
    from docx.shared import Pt, RGBColor

    code_text = _sanitize_xml_text(code_text)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _add_table_from_html(doc, table_element) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, RGBColor

    BLACK = RGBColor(0, 0, 0)
    rows_el = table_element.find_all("tr")
    if not rows_el:
        return
    first_row = rows_el[0]
    cols = first_row.find_all(["th", "td"])
    num_cols = len(cols)
    if num_cols == 0:
        return

    table = doc.add_table(rows=len(rows_el), cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows_el):
        cells = row.find_all(["th", "td"])
        for col_idx, cell in enumerate(cells):
            if col_idx < num_cols:
                table_cell = table.rows[row_idx].cells[col_idx]
                table_cell.text = _sanitize_xml_text(cell.get_text(strip=True))
                for paragraph in table_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BLACK
                        if cell.name == "th":
                            run.bold = True

    def set_cell_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            edge_el = OxmlElement(f"w:{edge}")
            edge_el.set(qn("w:val"), "single")
            edge_el.set(qn("w:sz"), "4")
            edge_el.set(qn("w:space"), "0")
            edge_el.set(qn("w:color"), "000000")
            tcBorders.append(edge_el)
        tcPr.append(tcBorders)

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)

    # Explicit column widths tuned to reduce aggressive wrapping in Word.
    if num_cols == 8:
        # Parameters Reference Table: #, Parameter, Description, Type, Length, Decimal, Data Element, Domain
        widths_8 = (0.5, 1.8, 2.7, 0.85, 0.75, 0.75, 1.3, 1.05)
        for i, w in enumerate(widths_8):
            if i < num_cols:
                table.columns[i].width = Inches(w)
    elif num_cols == 5:
        # EI Function Structure: Structure Name, Field Name, Description, Data Type, Component Type
        widths_5 = (1.5, 1.25, 2.5, 0.85, 1.1)
        for i, w in enumerate(widths_5):
            if i < num_cols:
                table.columns[i].width = Inches(w)

    doc.add_paragraph()


def _process_inline_formatting(paragraph, element, *, strip_dash: bool = False) -> None:
    from bs4 import NavigableString
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Pt, RGBColor

    BLACK = RGBColor(0, 0, 0)
    first_text = True

    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if strip_dash and first_text and text.lstrip().startswith("- "):
                text = text.lstrip().replace("- ", "", 1)
                first_text = False
            if text.strip():
                run = paragraph.add_run(text)
                run.font.color.rgb = BLACK
        elif child.name in ("strong", "b"):
            text = child.get_text()
            if strip_dash and first_text and text.startswith("- "):
                text = text.replace("- ", "", 1)
                first_text = False
            run = paragraph.add_run(text)
            run.bold = True
            run.font.color.rgb = BLACK
        elif child.name in ("em", "i"):
            run = paragraph.add_run(child.get_text())
            run.italic = True
            run.font.color.rgb = BLACK
        elif child.name == "code":
            run = paragraph.add_run(child.get_text())
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
        elif child.name == "a":
            run = paragraph.add_run(child.get_text())
            run.underline = True
            run.font.color.rgb = BLACK
        elif child.name == "mark":
            # Pipeline 04: dictionary explanation not yet in checked params.txt — yellow highlight in Word.
            txt = child.get_text()
            if txt.strip():
                run = paragraph.add_run(txt)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.color.rgb = BLACK
        else:
            _process_inline_formatting(paragraph, child, strip_dash=strip_dash and first_text)
            first_text = False


def _add_blank_paragraph(doc) -> None:
    """Insert one empty paragraph (one blank visual row in Word)."""
    doc.add_paragraph()


def _is_param_title_p(element) -> bool:
    """`**NAME** (Role):` row that starts a parameter block in section 04-style docs."""
    if element.name != "p":
        return False
    strong_tags = element.find_all("strong")
    if not strong_tags or len(strong_tags) != 1:
        return False
    first_strong = strong_tags[0]
    full_text = element.get_text()
    bold_text = first_strong.get_text()
    if not full_text.strip().startswith(bold_text.strip()):
        return False
    remainder = full_text[len(bold_text) :].strip()
    return remainder.startswith("(") and ":" in remainder


def _is_whole_paragraph_bold_single_strong(element) -> bool:
    if element.name != "p":
        return False
    strong_tags = element.find_all("strong")
    if not strong_tags or len(strong_tags) != 1:
        return False
    full_text = element.get_text(strip=True)
    bold_text = strong_tags[0].get_text(strip=True)
    return full_text.rstrip(":").strip() == bold_text.rstrip(":").strip()


def _iter_top_level_paragraphs(doc):
    """Body-level paragraphs only (excludes table cell paragraphs)."""
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)


def verify_docx_spacing(doc, *, md_path: Path | None = None) -> list[str]:
    """
    Verify one empty top-level paragraph before each section header and that the
    document contains at least one pair of consecutive empty rows.
    Parameter blocks are not required to have trailing blank rows.
    """
    errors: list[str] = []
    paras = list(_iter_top_level_paragraphs(doc))
    has_double_empty = False
    for i in range(len(paras) - 1):
        if not paras[i].text.strip() and not paras[i + 1].text.strip():
            has_double_empty = True
            break
    if not has_double_empty:
        errors.append("Expected at least one pair of consecutive empty rows, but none were found.")

    for i, p in enumerate(paras):
        st = p.style.name if p.style else ""
        if st not in ("Heading 2", "Heading 3"):
            continue
        if i == 0:
            continue
        prev_text = paras[i - 1].text.strip()
        if prev_text:
            errors.append(
                f"Section header {p.text!r}: expected one empty row immediately before it; "
                f"got {prev_text!r}"
            )

    # No enforced blank after each parameter block (requested behavior).

    return errors


def _render_paragraph_element(
    doc, element, *, current_section: str, no_bullet_sections: list[str], BLACK
) -> str:
    """
    Render a single BeautifulSoup <p> to docx. Returns param_state hint:
    "param_title", "whole_bold", "param_bullets_done", or "".
    """
    strong_tags = element.find_all("strong")
    if strong_tags and len(strong_tags) == 1:
        full_text_strip = element.get_text(strip=True)
        bold_text = strong_tags[0].get_text(strip=True)
        if full_text_strip.rstrip(":").strip() == bold_text.rstrip(":").strip():
            p = doc.add_paragraph()
            run = p.add_run(element.get_text(strip=True))
            run.bold = True
            run.font.color.rgb = BLACK
            return "whole_bold"

        first_strong = strong_tags[0]
        full_text = element.get_text()
        bold_text = first_strong.get_text()
        if full_text.strip().startswith(bold_text.strip()):
            remainder = full_text[len(bold_text) :].strip()
            if remainder.startswith("(") and ":" in remainder:
                p = doc.add_paragraph()
                run = p.add_run(bold_text.strip())
                run.bold = True
                run.font.color.rgb = BLACK
                p.add_run(" " + remainder).font.color.rgb = BLACK
                return "param_title"
            if not remainder or remainder.startswith("-") or remainder.startswith("•"):
                p = doc.add_paragraph()
                run = p.add_run(bold_text.strip())
                run.bold = True
                run.font.color.rgb = BLACK
                if remainder and (remainder.startswith("-") or remainder.startswith("•")):
                    for line in remainder.split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        text = line[2:] if (line.startswith("- ") or line.startswith("• ")) else line
                        if _md_docx_no_list_styles(current_section, no_bullet_sections, options_ul=False):
                            p = doc.add_paragraph()
                            p.add_run(text).font.color.rgb = BLACK
                        else:
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(text).font.color.rgb = BLACK
                return "param_bullets_done"

    text_content = element.get_text()
    if "\n- " in text_content or text_content.strip().startswith("- "):
        for line in text_content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("- "):
                text = line[2:]
                if _md_docx_no_list_styles(current_section, no_bullet_sections, options_ul=False):
                    p = doc.add_paragraph()
                    p.add_run(text).font.color.rgb = BLACK
                else:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(text).font.color.rgb = BLACK
            else:
                p = doc.add_paragraph()
                p.add_run(line).font.color.rgb = BLACK
        return ""

    p = doc.add_paragraph()
    _process_inline_formatting(p, element)
    return ""


def convert_md_to_docx(
    md_path: Path | str,
    output_path: Path | str | None = None,
    *,
    skip_spacing_verify: bool = False,
) -> Path:
    """
    Convert a markdown file to a Word document using python-docx (no Pandoc).
    Requires: python-docx, markdown, beautifulsoup4.

    skip_spacing_verify: when True, skip verify_docx_spacing (for auxiliary exports
    that do not follow the full EI pipeline markdown layout).
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    md_path = Path(md_path)
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {md_path}")

    out = Path(output_path) if output_path is not None else md_path.with_suffix(".docx")

    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()
    md_content = _sanitize_xml_text(md_content)

    doc = Document()
    BLACK = RGBColor(0, 0, 0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK

    for i in range(1, 10):
        try:
            doc.styles[f"Heading {i}"].font.color.rgb = BLACK
        except KeyError:
            pass

    soup = _parse_markdown_to_elements(md_content)
    no_bullet_sections = ["problem description", "suggested resolution"]
    current_section = ""
    elements = [c for c in soup.children if getattr(c, "name", None)]
    idx = 0
    # idle | want_body | want_options_line | want_ul
    param_state = "idle"

    def _close_open_param_block() -> None:
        nonlocal param_state
        if param_state != "idle":
            param_state = "idle"

    while idx < len(elements):
        element = elements[idx]

        if param_state == "want_options_line" and element.name != "p":
            param_state = "idle"
            continue
        if param_state == "want_ul" and element.name != "ul":
            param_state = "idle"
            continue

        if element.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            _close_open_param_block()
            if list(_iter_top_level_paragraphs(doc)):
                _add_blank_paragraph(doc)
            level = int(element.name[1])
            heading_text = _sanitize_xml_text(element.get_text(strip=True))
            if heading_text.strip().lower() == "suggested resolution":
                level = 2
            doc.add_heading(heading_text, level=min(level, 4))
            if level in (2, 3):
                current_section = heading_text.lower().strip()
            idx += 1
            continue

        if element.name in ("table", "pre"):
            _close_open_param_block()

        if element.name == "p" and _is_param_title_p(element):
            _close_open_param_block()
            _render_paragraph_element(
                doc, element, current_section=current_section, no_bullet_sections=no_bullet_sections, BLACK=BLACK
            )
            param_state = "want_body"
            idx += 1
            continue

        if element.name == "p" and param_state == "want_options_line":
            _render_paragraph_element(
                doc, element, current_section=current_section, no_bullet_sections=no_bullet_sections, BLACK=BLACK
            )
            param_state = "want_ul"
            idx += 1
            continue

        if element.name == "p" and param_state == "want_body":
            if _is_param_title_p(element):
                param_state = "idle"
                _render_paragraph_element(
                    doc, element, current_section=current_section, no_bullet_sections=no_bullet_sections, BLACK=BLACK
                )
                param_state = "want_body"
                idx += 1
                continue

            hint = _render_paragraph_element(
                doc, element, current_section=current_section, no_bullet_sections=no_bullet_sections, BLACK=BLACK
            )
            nxt = elements[idx + 1] if idx + 1 < len(elements) else None
            if hint == "param_bullets_done":
                param_state = "idle"
            elif hint == "whole_bold":
                if "Options" in element.get_text():
                    param_state = "want_ul"
                elif (
                    nxt
                    and nxt.name == "p"
                    and _is_whole_paragraph_bold_single_strong(nxt)
                    and "Options" in nxt.get_text()
                ):
                    param_state = "want_options_line"
                else:
                    param_state = "idle"
            else:
                if (
                    nxt
                    and nxt.name == "p"
                    and _is_whole_paragraph_bold_single_strong(nxt)
                    and "Options" in nxt.get_text()
                ):
                    param_state = "want_options_line"
                elif nxt and nxt.name == "ul":
                    param_state = "want_ul"
                else:
                    param_state = "idle"
            idx += 1
            continue

        if element.name == "ul" and param_state == "want_ul":
            for li in element.find_all("li", recursive=False):
                if _md_docx_no_list_styles(current_section, no_bullet_sections, options_ul=True):
                    p = doc.add_paragraph()
                else:
                    p = doc.add_paragraph(style="List Bullet")
                _process_inline_formatting(p, li)
            param_state = "idle"
            idx += 1
            continue

        if element.name == "ul" and param_state == "want_body":
            for li in element.find_all("li", recursive=False):
                if _md_docx_no_list_styles(current_section, no_bullet_sections, options_ul=False):
                    p = doc.add_paragraph()
                else:
                    p = doc.add_paragraph(style="List Bullet")
                _process_inline_formatting(p, li)
            param_state = "idle"
            idx += 1
            continue

        if element.name == "p" and param_state == "idle":
            hint = _render_paragraph_element(
                doc, element, current_section=current_section, no_bullet_sections=no_bullet_sections, BLACK=BLACK
            )
            idx += 1
            continue

        if element.name == "ul":
            for li in element.find_all("li", recursive=False):
                if _md_docx_no_list_styles(current_section, no_bullet_sections, options_ul=False):
                    p = doc.add_paragraph()
                else:
                    p = doc.add_paragraph(style="List Bullet")
                _process_inline_formatting(p, li)

        elif element.name == "ol":
            if param_state == "want_body":
                for i, li in enumerate(element.find_all("li", recursive=False), start=1):
                    if _md_docx_no_list_styles(current_section, no_bullet_sections, options_ul=False):
                        p = doc.add_paragraph()
                        r = p.add_run(f"{i}. ")
                        r.bold = False
                        r.font.color.rgb = BLACK
                        _process_inline_formatting(p, li)
                    else:
                        p = doc.add_paragraph(style="List Number")
                        _process_inline_formatting(p, li)
                param_state = "idle"
            else:
                for i, li in enumerate(element.find_all("li", recursive=False), start=1):
                    if _md_docx_no_list_styles(current_section, no_bullet_sections, options_ul=False):
                        p = doc.add_paragraph()
                        r = p.add_run(f"{i}. ")
                        r.font.color.rgb = BLACK
                        _process_inline_formatting(p, li)
                    else:
                        p = doc.add_paragraph(style="List Number")
                        _process_inline_formatting(p, li)

        elif element.name == "table":
            _add_table_from_html(doc, element)

        elif element.name == "pre":
            code = element.find("code")
            _add_code_block(doc, code.get_text() if code else element.get_text())

        elif element.name == "hr":
            pass

        idx += 1

    _close_open_param_block()

    if not skip_spacing_verify:
        spacing_errors = verify_docx_spacing(doc, md_path=md_path)
        if spacing_errors:
            raise ValueError("DOCX spacing verification failed:\n- " + "\n- ".join(spacing_errors))

    doc.save(str(out))
    return out
