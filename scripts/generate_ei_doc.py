#!/usr/bin/env python3
"""
Generate EI (Exception Indicator) documentation from input/ and reference/.
Option C: full Python workflow — extract, resolve params, fill sections, emit to output/.
All paths relative to project root (parent of scripts/). Run from project root.

Narrative sections (Overview, Problem, Resolution) are LLM-generated when --no-llm is not passed
and one of these is set (checked in order): OPENAI_API_KEY, ANTHROPIC_API_KEY,
or EI_DOC_LLM_API_KEY + EI_DOC_LLM_BASE_URL (OpenAI-compatible). No API keys in source.

With --cursor-workflow: output skeleton with [TODO: General Overview], [TODO: Problem Description],
[TODO: Suggested Resolution] and write Explanation_<base>_narrative_context.md; do not call LLM
or use built-in narrative for those three sections. Cursor fills them using the context file.
"""
from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

# Project root = parent of directory containing this script
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent


def _ensure_deps():
    try:
        import openpyxl  # noqa: F401
        from docx import Document  # noqa: F401
    except ImportError:
        print("Missing deps. Install with: pip install -r scripts/requirements.txt", file=sys.stderr)
        sys.exit(1)


def discover_inputs(input_dir: Path):
    """Find one Code_*.txt, one Structure_*.xlsx, one Available fields_*.xlsx."""
    code_files = list(input_dir.glob("Code_*.txt"))
    struct_files = list(input_dir.glob("Structure_*.xlsx"))
    avail_files = list(input_dir.glob("Available fields_*.xlsx"))
    if not code_files or not struct_files or not avail_files:
        raise SystemExit(
            "Need exactly one of each in input/: Code_*.txt, Structure_*.xlsx, Available fields_*.xlsx"
        )
    if len(code_files) > 1 or len(struct_files) > 1 or len(avail_files) > 1:
        raise SystemExit("Multiple matches for Code/Structure/Available fields; use one function set per input/.")
    return code_files[0], struct_files[0], avail_files[0]


def load_parameters(avail_path: Path) -> list[dict]:
    """Load Parameters sheet: rows from row 3, headers in row 2. Column 'Field' -> Parameter."""
    import openpyxl
    wb = openpyxl.load_workbook(avail_path, read_only=True)
    if "Parameters" not in wb.sheetnames:
        raise SystemExit(f"No 'Parameters' sheet in {avail_path}")
    ws = wb["Parameters"]
    rows = list(ws.iter_rows(min_row=1, max_row=200, values_only=True))
    wb.close()
    # Row 1 may be title; row 2 = headers
    header_row = 1
    for i, r in enumerate(rows):
        if r and str(r[0]).strip().lower() in ("field", "parameter"):
            header_row = i + 1
            break
    headers = [str(c).strip() if c is not None else "" for c in rows[header_row - 1]]
    # Normalize: first column often "Field" -> we output as Parameter
    param_col = "Field" if "Field" in headers else (headers[0] if headers else "Parameter")
    col_map = {h or f"__{i}": i for i, h in enumerate(headers)}
    out = []
    for row in rows[header_row:]:
        if not any(c is not None for c in row):
            continue
        vals = [row[i] if i < len(row) else None for i in range(max(col_map.values()) + 1)]
        param = str(vals[col_map.get(param_col, 0)] or "").strip()
        if not param:
            continue
        out.append({
            "Parameter": param,
            "Description": _cell(vals, col_map, "Description"),
            "Type": _cell(vals, col_map, "Type"),
            "Length": _cell(vals, col_map, "Length"),
            "Decimal": _cell(vals, col_map, "Decimal"),
            "Data Element": _cell(vals, col_map, "Data Element"),
            "Domain": _cell(vals, col_map, "Domain"),
        })
    return out


def _cell(vals: list, col_map: dict, name: str) -> str:
    idx = col_map.get(name, -1)
    if idx < 0:
        return ""
    v = vals[idx] if idx < len(vals) else None
    return "" if v is None else str(v).strip()


def load_structure(struct_path: Path) -> list[dict]:
    """Load first sheet: row 1 = headers, data from row 2. Columns: Structure Name, Field Name, Description, Data Type, Component Type."""
    import openpyxl
    wb = openpyxl.load_workbook(struct_path, read_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(min_row=1, max_row=200, values_only=True))
    wb.close()
    if not rows:
        return []
    headers = [str(c).strip() if c is not None else "" for c in rows[0]]
    col_map = {h or f"__{i}": i for i, h in enumerate(headers)}
    names = ["Structure Name", "Field Name", "Description", "Data Type", "Component Type"]
    out = []
    for row in rows[1:]:
        if not any(c is not None for c in row):
            continue
        vals = list(row) if isinstance(row, (list, tuple)) else [row]
        rec = {}
        for n in names:
            idx = next((col_map[k] for k in col_map if n.lower() in k.lower() or k == n), None)
            if idx is None:
                idx = names.index(n) if n in names else -1
            rec[n] = _cell(vals, {n: names.index(n)} if n in names else col_map, n) if n in col_map else (str(vals[names.index(n)]).strip() if names.index(n) < len(vals) else "")
            for k in col_map:
                if k.strip() == n or (n.lower() in k.lower() and "structure" in n.lower() and "structure" in k.lower()) or (n.lower() in k.lower() and "field" in n.lower() and "field" in k.lower()):
                    rec[n] = str(vals[col_map[k]]).strip() if col_map[k] < len(vals) else ""
                    break
        # Fallback: by position
        if not rec:
            rec = {
                "Structure Name": str(vals[0]) if len(vals) > 0 else "",
                "Field Name": str(vals[1]) if len(vals) > 1 else "",
                "Description": str(vals[2]) if len(vals) > 2 else "",
                "Data Type": str(vals[3]) if len(vals) > 3 else "",
                "Component Type": str(vals[4]) if len(vals) > 4 else "",
            }
        else:
            for i, n in enumerate(names):
                if n not in rec or rec[n] == "":
                    rec[n] = str(vals[i]) if i < len(vals) else ""
        out.append(rec)
    return out


def load_structure_simple(struct_path: Path) -> list[dict]:
    """Load structure by position: row 0 = headers, row 1+ = data."""
    import openpyxl
    wb = openpyxl.load_workbook(struct_path, read_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(min_row=1, max_row=200, values_only=True))
    wb.close()
    if len(rows) < 2:
        return []
    h = [str(x).strip() if x is not None else "" for x in rows[0]]
    names = ["Structure Name", "Field Name", "Description", "Data Type", "Component Type"]
    out = []
    for r in rows[1:]:
        row = list(r) if r else []
        while len(row) < 5:
            row.append(None)
        out.append({
            names[0]: str(row[0]) if row[0] is not None else "",
            names[1]: str(row[1]) if row[1] is not None else "",
            names[2]: str(row[2]) if row[2] is not None else "",
            names[3]: str(row[3]) if row[3] is not None else "",
            names[4]: str(row[4]) if row[4] is not None else "",
        })
    return out


def load_param_reference(docx_path: Path, param_names: set[str]) -> dict[str, str]:
    """Build param_name -> text from _Parameters.docx. Collect paragraphs that contain param name."""
    from docx import Document
    doc = Document(docx_path)
    result = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        for name in param_names:
            if name not in result and re.search(r"\b" + re.escape(name) + r"\b", t, re.I):
                # Use this paragraph as seed; optionally concatenate following paras
                result[name] = t
                break
    # For each param still missing, use generic
    for name in param_names:
        if name not in result:
            result[name] = f"Configuration guidance for {name}. Refer to parameter reference or code analysis."
    return result


def load_param_reference_docx(docx_path: Path, param_names: set[str]) -> dict[str, str]:
    from docx import Document
    doc = Document(docx_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    result = {}
    for name in param_names:
        # Collect sentences/paragraphs that contain this param name
        pattern = r"([^\n]*(?:\b" + re.escape(name) + r"\b)[^\n]*)"
        hits = re.findall(pattern, all_text, re.I)
        if hits:
            result[name] = " ".join(hits).strip() or f"Configuration for {name}. See parameter reference."
        else:
            result[name] = f"Configuration for {name}. See parameter reference or code analysis."
    return result


def extract_defaults(code_text: str) -> dict[str, str]:
    """Extract lv_* = '...' or lv_* = number. Map lv_NAME -> NAME (uppercase). Only literal RHS."""
    defaults = {}
    # lv_DATE_REF_FLD = 'VDATU'  or  lv_backdays = 0  (skip lv_backdays = lv_forwdays * ...)
    for m in re.finditer(r"lv_(\w+)\s*=\s*('([^']*)'|\"([^\"]*)\"|(\d+))", code_text):
        param = m.group(1).upper()
        val = m.group(3) or m.group(4) or m.group(5) or ""
        if not val or val.startswith("lv_") or "*" in val:
            continue
        defaults[param] = val
    # map common ABAP-style to param name (e.g. BACKDAYS from lv_backdays)
    name_map = {"BACKDAYS": "BACKDAYS", "AGGR_PERIOD": "AGGR_PERIOD", "DATE_REF_FLD": "DATE_REF_FLD", "FORWDAYS": "FORWDAYS"}
    out = {}
    for k, v in defaults.items():
        key = name_map.get(k, k)
        if key not in out:
            out[key] = v
    return out


def extract_technical_code(code_text: str) -> str:
    """FUNCTION /SKN/F_SW_10_01_ORD_VAL_TOT -> SW_10_01_ORD_VAL_TOT."""
    m = re.search(r"FUNCTION\s+(?:/\w+/)?F?_?(\w+)", code_text, re.I)
    return m.group(1) if m else ""


def full_name_from_available_path(avail_path: Path, technical_code: str) -> str:
    """From 'Available fields_Exceptional sales documents values – Aggregated_SW_10_01_ORD_VAL_TOT.xlsx' -> Exceptional sales documents values – Aggregated."""
    stem = avail_path.stem
    prefix = "Available fields_"
    if stem.startswith(prefix):
        stem = stem[len(prefix):]
    if technical_code and stem.endswith("_" + technical_code):
        stem = stem[: -len(technical_code) - 1]
    return stem.replace("_", " ").strip()


def render_doc(
    *,
    full_name: str,
    technical_code: str,
    params: list[dict],
    param_texts: dict[str, str],
    structure_rows: list[dict],
    code_text: str,
    defaults: dict[str, str],
    sample_format_path: Path,
    overview: str | None = None,
    problem: str | None = None,
    resolution: str | None = None,
) -> str:
    """Build markdown following sample_format section order and table shapes."""
    with open(sample_format_path, "r", encoding="utf-8") as f:
        sample = f.read()

    # Table: Parameters (7 cols)
    param_header = "| Parameter | Description | Type | Length | Decimal | Data Element | Domain |"
    param_sep = "|-----------|-------------|------|--------|---------|--------------|--------|"
    param_rows = []
    for p in params:
        cells = [
            str(p.get("Parameter", "") or ""),
            str(p.get("Description", "") or ""),
            str(p.get("Type", "") or ""),
            str(p.get("Length", "") or ""),
            str(p.get("Decimal", "") or ""),
            str(p.get("Data Element", "") or ""),
            str(p.get("Domain", "") or ""),
        ]
        param_rows.append("| " + " | ".join(cells) + " |")
    params_table = "\n".join([param_header, param_sep] + param_rows)

    # Parameter Configuration Guidelines: **PARAM** (Description): text ---
    guideline_lines = []
    for p in params:
        name = p.get("Parameter", "")
        desc = p.get("Description", "")
        text = param_texts.get(name, param_texts.get(name.upper(), f"Configuration for {name}. See parameter reference."))
        guideline_lines.append(f"**{name}** ({desc}):\n\n{text}\n\n")

    # Default Values (only those in code)
    default_lines = [f"- **{k}** — Default: `{v}`" for k, v in defaults.items()]

    # Structure table (5 cols)
    struct_header = "| Structure Name | Field Name | Description | Data Type | Component Type |"
    struct_sep = "|----------------|------------|-------------|-----------|----------------|"
    struct_rows = []
    for s in structure_rows:
        cells = [
            str(s.get("Structure Name", "") or ""),
            str(s.get("Field Name", "") or ""),
            str(s.get("Description", "") or ""),
            str(s.get("Data Type", "") or ""),
            str(s.get("Component Type", "") or ""),
        ]
        struct_rows.append("| " + " | ".join(cells) + " |")
    structure_table = "\n".join([struct_header, struct_sep] + struct_rows)

    # Section order from sample_format
    h1 = f"# Exception Indicator: {full_name.title()} ({technical_code})"
    # Richer built-in fallback (benchmark-style structure) when LLM is not used
    _overview_fallback = (
        "This Exception Indicator (EI) monitors aggregated sales order values in Sales and Distribution (SD) to identify exceptional patterns in sales document values across configurable time periods and organizational dimensions. It provides consolidated visibility into sales volume trends, enabling detection of unusual value concentrations, high-value transactions, and sales pattern anomalies that require management attention.\n\n"
        "This EI serves as an essential control for sales management and financial oversight by:\n"
        "- Aggregating sales order net values by customizable time periods (Month/Week/Quarter/Year) and organizational dimensions\n"
        "- Identifying exceptional sales volumes that exceed predefined thresholds for transaction counts or value totals\n"
        "- Monitoring sales patterns across multiple currencies (document currency and foreign currency)\n"
        "- Providing flexible filtering by sales organization, distribution channel, division, customer, and document characteristics\n"
        "- Supporting multi-partner analysis through three configurable business partner fields for comprehensive sales relationship visibility\n\n"
        "This aggregated monitoring enables organizations to detect high-value transaction clusters, unusual sales concentration patterns, potential revenue recognition issues, and sales trends requiring executive visibility. The EI is particularly valuable for month-end close processes, sales performance reviews, and financial exception management.\n\n"
        "The EI retrieves detailed sales order data from SAP SD tables (VBAK - Sales Document Header, VBPA - Sales Document Partner, KNA1 - Customer Master), then aggregates the results by the selected time period and organizational dimensions. It calculates total counts (TOT_CNT), total document currency values (TOT_NETWR), and total foreign currency values (TOT_NETWR_FR) for each aggregation bucket, then filters based on user-specified threshold ranges."
    )
    _problem_fallback = (
        "Failure to monitor aggregated sales order values creates multiple risks across financial reporting, operational management, and compliance:\n\n"
        "**Financial and Reporting Issues**\n"
        "- Undetected high-value order concentrations can distort period-over-period sales trend analysis and forecasting accuracy\n"
        "- Exceptional sales volumes in specific periods may indicate revenue recognition timing issues or premature booking\n"
        "- Unusual value patterns in foreign currency transactions can signal currency risk exposure requiring hedging actions\n"
        "- Aggregated sales anomalies may delay month-end close processes when discovered late during financial review\n"
        "- Concentrated high-value orders in specific sales organizations or channels can mask underlying performance issues in other business units\n\n"
        "**Sales Operations and Control Risks**\n"
        "- Large transaction clusters without proper visibility may indicate unauthorized discounting or pricing violations\n"
        "- Exceptional values in specific customer segments could signal credit risk concentration requiring management intervention\n"
        "- Unusual sales patterns by partner functions (sold-to, ship-to, bill-to) may indicate customer master data quality issues\n"
        "- High-volume activity in specific divisions or distribution channels may reflect operational bottlenecks or resource constraints\n"
        "- Atypical aggregated values could indicate data entry errors or system integration failures requiring immediate correction\n\n"
        "**Management Visibility and Decision-Making Risks**\n"
        "- Lack of aggregated value monitoring delays executive awareness of significant business trends and market shifts\n"
        "- Unidentified sales concentration patterns can lead to missed opportunities for strategic pricing or customer engagement\n"
        "- Exceptional transaction volumes may require additional audit scrutiny or compliance review but go unnoticed\n"
        "- Absence of multi-dimensional sales analysis limits ability to optimize sales territory assignments and resource allocation"
    )
    _resolution_fallback = (
        "**Immediate Response**\n"
        "- Review the aggregated sales values flagged by the EI to understand the nature and scope of the exceptional pattern (threshold violations, value concentration, period-specific spikes)\n"
        "- Verify the authenticity of high-value orders using transaction VA03 (Display Sales Order) to confirm legitimacy and proper authorization\n"
        "- Check sales document status and processing progress to ensure no manual intervention or corrections are pending\n"
        "- Identify the business context for exceptional volumes: promotional campaigns, large customer orders, seasonal patterns, or data quality issues\n\n"
        "**System Assessment**\n"
        "- Analyze the aggregation dimensions (time period, organizational fields) to understand which factors drive the exceptional pattern\n"
        "- Review historical trends by comparing current aggregated values to prior periods using the same aggregation criteria\n"
        "- Examine currency-specific patterns by comparing TOT_NETWR (document currency) and TOT_NETWR_FR (foreign currency) totals\n"
        "- Assess partner function distribution (BP1, BP2, BP3) to identify customer relationship patterns or master data inconsistencies\n"
        "- Investigate sales document characteristics (AUART - document type, VBTYP - document category) to determine if exceptions are type-specific\n"
        "- Validate date reference field usage to ensure appropriate timing basis for aggregation (VDATU - requested delivery date, AUDAT - document date, ERDAT - creation date)\n\n"
        "**Corrective Actions**\n"
        "- If unauthorized or erroneous orders are identified, initiate sales document correction procedures using VA02 (Change Sales Order)\n"
        "- For legitimate high-value orders requiring special approval, escalate to sales management and finance for validation\n"
        "- Update customer master data (VD02 - Change Customer Sales Data) if partner function issues or credit limit violations are detected\n"
        "- Adjust pricing or discounting arrangements using VK11 (Create Condition Record) if pricing violations are confirmed\n"
        "- Implement additional monitoring controls by adjusting EI parameters to tighten threshold criteria for future executions\n"
        "- Document exceptional patterns and business justifications for audit trail and management reporting purposes\n"
        "- Establish recurring EI execution schedules to provide continuous visibility into sales value trends and concentration risks\n"
        "- Configure Dynamic Recipient List (USER_FLD parameter) to automatically route alerts to appropriate sales managers and finance stakeholders based on organizational responsibility"
    )
    overview = overview if overview is not None else _overview_fallback
    problem = problem if problem is not None else _problem_fallback
    resolution = resolution if resolution is not None else _resolution_fallback
    param_rels = "Parameters work together to define filters and aggregation. BACKDAYS and date reference define the time window; aggregation and field parameters control grouping."
    practical = "```\nBACKDAYS=7, DATE_REF_FLD=VDATU, AGGR_PERIOD=M — Last 7 days, monthly by requested delivery date\n```"

    sections = [
        h1,
        "",
        "## General Overview",
        "",
        overview,
        "",
        "---",
        "",
        "## Problem Description",
        "",
        problem,
        "",
        "---",
        "",
        "## Suggested Resolution",
        "",
        resolution,
        "",
        "---",
        "",
        "## Parameters",
        "",
        "### Parameters Reference Table",
        "",
        "This table lists all configurable input parameters. Users set values for these parameters to filter and control the EI's data selection and processing logic.",
        "",
        params_table,
        "",
        "### Parameter Configuration Guidelines",
        "",
        "**IMPORTANT:** This section provides configuration guidance for ALL parameters listed in the Parameters Reference Table above.",
        "",
        "\n\n".join(guideline_lines),
        "### Parameter Relationships",
        "",
        param_rels,
        "",
        "### Default Values and Parameter Options Explicitly Stated in EI Code",
        "",
        "\n".join(default_lines) if default_lines else "- (None found in code.)",
        "",
        "### Practical Configuration Examples",
        "",
        practical,
        "",
        "## EI Function Structure",
        "",
        "This table lists all output fields returned by the EI. These fields contain the results of the EI's data retrieval and calculations.",
        "",
        structure_table,
        "",
        "## ABAP Code",
        "",
        "```abap",
        code_text.strip(),
        "```",
    ]
    return "\n".join(sections)


def main() -> None:
    _ensure_deps()
    parser = argparse.ArgumentParser(description="Generate EI doc from input/ and reference/")
    parser.add_argument("--input", default="input", help="Input folder (per-function)")
    parser.add_argument("--reference", default="reference", help="Reference folder (sample_format, _Parameters)")
    parser.add_argument("--output", default="output", help="Output folder")
    parser.add_argument("--project-root", type=Path, default=PROJECT_ROOT, help="Project root")
    parser.add_argument("--no-llm", action="store_true", help="Skip LLM; use built-in narrative templates. Otherwise uses first available of: OPENAI_API_KEY, ANTHROPIC_API_KEY, or EI_DOC_LLM_API_KEY+EI_DOC_LLM_BASE_URL.")
    parser.add_argument("--cursor-workflow", action="store_true", help="Output skeleton with [TODO:] for Overview/Problem/Resolution and write narrative_context.md; do not call LLM. Cursor fills TODOs using that context and .cursor rules.")
    args = parser.parse_args()
    root = Path(args.project_root)
    input_dir = root / args.input
    ref_dir = root / args.reference
    out_dir = root / args.output

    if not input_dir.is_dir():
        raise SystemExit(f"Input folder not found: {input_dir}")
    if not ref_dir.is_dir():
        raise SystemExit(f"Reference folder not found: {ref_dir}")
    out_dir.mkdir(parents=True, exist_ok=True)

    code_path, struct_path, avail_path = discover_inputs(input_dir)
    base_name = code_path.stem
    if base_name.upper().startswith("CODE_"):
        base_name = base_name[5:]
    elif base_name.upper().startswith("CODE"):
        base_name = base_name[4:].lstrip("_")

    sample_format = ref_dir / "sample_format.md"
    param_ref_docx = ref_dir / "_Parameters.docx"
    if not sample_format.exists():
        raise SystemExit(f"Sample format not found: {sample_format}")
    if not param_ref_docx.exists():
        raise SystemExit(f"Parameter reference not found: {param_ref_docx}")

    code_text = code_path.read_text(encoding="utf-8", errors="replace")
    params = load_parameters(avail_path)
    structure_rows = load_structure_simple(struct_path)
    param_names = {p["Parameter"] for p in params}
    param_texts = load_param_reference_docx(param_ref_docx, param_names)
    defaults = extract_defaults(code_text)
    technical_code = extract_technical_code(code_text)
    full_name = full_name_from_available_path(avail_path, technical_code)

    overview, problem, resolution = None, None, None
    env = os.environ

    def _llm_ready() -> tuple[str, str, str | None]:
        """Return (provider, api_key, base_url or None). Empty key means not ready."""
        if env.get("OPENAI_API_KEY", "").strip():
            return "openai", env["OPENAI_API_KEY"].strip(), env.get("EI_DOC_LLM_BASE_URL", "").strip() or None
        if env.get("ANTHROPIC_API_KEY", "").strip():
            return "anthropic", env["ANTHROPIC_API_KEY"].strip(), None
        key, base = env.get("EI_DOC_LLM_API_KEY", "").strip(), env.get("EI_DOC_LLM_BASE_URL", "").strip()
        if key and base:
            return "openai", key, base.rstrip("/")
        return "", "", None

    if args.cursor_workflow:
        overview = "[TODO: General Overview]"
        problem = "[TODO: Problem Description]"
        resolution = "[TODO: Suggested Resolution]"
    elif args.no_llm:
        print("Narrative: using built-in templates (--no-llm).", file=sys.stderr)
    else:
        provider, api_key, base_url = _llm_ready()
        if not api_key:
            print(
                "Narrative: using built-in templates (no LLM key). Set one of: OPENAI_API_KEY, ANTHROPIC_API_KEY, or EI_DOC_LLM_API_KEY+EI_DOC_LLM_BASE_URL.",
                file=sys.stderr,
            )
        else:
            structure_summary = ", ".join(
                str(s.get("Field Name", "") or "").strip() or str(s.get("Description", "") or "").strip()
                for s in structure_rows[:50]
            ).strip() or "Output fields from structure."
            parameters_summary = [f"{p.get('Parameter', '')} ({p.get('Description', '') or ''})" for p in params]
            context = {
                "full_name": full_name,
                "technical_code": technical_code,
                "code_text": code_text,
                "structure_summary": structure_summary,
                "parameters_summary": parameters_summary,
            }
            try:
                if str(SCRIPT_DIR) not in sys.path:
                    sys.path.insert(0, str(SCRIPT_DIR))
                from llm_narrative import generate_narrative_sections
                narrative = generate_narrative_sections(
                    context,
                    api_key=api_key,
                    provider=provider,
                    base_url=base_url,
                )
                if narrative and (narrative.get("overview") or narrative.get("problem") or narrative.get("resolution")):
                    overview = narrative.get("overview") or None
                    problem = narrative.get("problem") or None
                    resolution = narrative.get("resolution") or None
                    print(f"Narrative: using LLM-generated Overview, Problem, Resolution ({provider}).", file=sys.stderr)
                else:
                    print("Narrative: using built-in templates (LLM returned no usable sections).", file=sys.stderr)
            except Exception as e:
                print(f"Narrative: using built-in templates (LLM failed: {e}).", file=sys.stderr)

    md = render_doc(
        full_name=full_name,
        technical_code=technical_code,
        params=params,
        param_texts=param_texts,
        structure_rows=structure_rows,
        code_text=code_text,
        defaults=defaults,
        sample_format_path=sample_format,
        overview=overview,
        problem=problem,
        resolution=resolution,
    )

    out_file = out_dir / f"Explanation_{base_name}.md"
    out_file.write_text(md, encoding="utf-8")
    print(f"Wrote {out_file}")

    if args.cursor_workflow:
        ref_files = root / "reference files"
        benchmark_candidates = list(ref_files.glob("Explanation_*.md")) if ref_files.is_dir() else []
        benchmark_path = str(benchmark_candidates[0]) if benchmark_candidates else "reference files/Explanation_*.md"
        structure_summary = ", ".join(
            str(s.get("Field Name", "") or "").strip() or str(s.get("Description", "") or "").strip()
            for s in structure_rows[:50]
        ).strip() or "Output fields from structure."
        parameters_summary = [f"{p.get('Parameter', '')} ({p.get('Description', '') or ''})" for p in params]
        code_excerpt = code_text[:8000] if len(code_text) > 8000 else code_text
        instructions = (
            "Use this context plus reference/sample_format.md and reference/_Parameters.docx to generate "
            "General Overview, Problem Description, Suggested Resolution. **Strictly follow all .cursor rules** "
            "(abap-function-doc-format-authority, document-authority, anti-hallucination-rules, behavior/verify-after-write). "
            "**Verify each drafted section** against (a) .cursor rules, (b) reference/sample_format.md, (c) benchmark (path below) "
            "**before** replacing the placeholder. Replace a [TODO: ...] only when the content passes all three. "
            "Read the file after each edit to confirm."
        )
        context_body = f"""# Narrative context for {out_file.name}

- **full_name:** {full_name}
- **technical_code:** {technical_code}

## Structure summary (output fields)
{structure_summary}

## Parameters (Name (Description))
"""
        for line in parameters_summary:
            context_body += f"- {line}\n"
        context_body += f"""
## Code path or excerpt
Path: {code_path}
Excerpt (first 8000 chars):
```
{code_excerpt}
```

## Benchmark path (for verification)
{benchmark_path}

## Instructions for Cursor
{instructions}
"""
        context_file = out_dir / f"Explanation_{base_name}_narrative_context.md"
        context_file.write_text(context_body, encoding="utf-8")
        print(f"Wrote {context_file}", file=sys.stderr)
        print(
            f'Next: Ask Cursor to "Fill all [TODO:] in output/{out_file.name} using output/{context_file.name} and reference/."',
            file=sys.stderr,
        )


if __name__ == "__main__":
    main()
