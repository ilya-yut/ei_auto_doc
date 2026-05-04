from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook


ROOT = Path(r"c:\vibe code dev\ei_auto_doc")
PART1_DIR = ROOT / "ei docs for analysis" / "EI docs" / "Part 1"
TEMP_DIR = ROOT / "ei docs for analysis" / "EI docs" / "temp"
OUT_DIR = TEMP_DIR / "Part1_enriched"

DICT_XLSX = TEMP_DIR / "params_dictionary.xlsx"
GE3_DOCX = ROOT / "ei docs for analysis" / "Part 1 conv" / "SHARED_PARAMETER_EXPLANATIONS_GE3_NO_BACKDAYS_USER_FLD.docx"

RE_PARAM_HEAD = re.compile(r"^([A-Z][A-Z0-9_]*)\s*\(.*\)\s*:?\s*$")
RE_OPTIONS = re.compile(r"^(?:\*?\*?)?[A-Z][A-Z0-9_]*\s*Options:\s*(?:\*?\*?)?$", re.I)
STOP_TITLES = {
    "parameter relationship",
    "parameter relationships",
    "default values",
    "practical example of parameter configuration",
    "practical configuration examples",
    "ei function structure",
    "abap code",
}


def insert_paragraph_after(paragraph: Paragraph, text: str, style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def load_dict_map(path: Path) -> dict[str, str]:
    wb = load_workbook(path, read_only=True)
    ws = wb["dictionary"] if "dictionary" in wb.sheetnames else wb[wb.sheetnames[0]]
    out: dict[str, str] = {}
    for r in range(2, ws.max_row + 1):
        p = ws.cell(r, 1).value
        e = ws.cell(r, 2).value
        if not p or not e:
            continue
        param = str(p).strip().upper()
        expl = str(e).strip()
        if param and expl:
            out[param] = expl
    return out


def load_ge3_map(path: Path) -> dict[str, str]:
    doc = Document(str(path))
    out: dict[str, str] = {}
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        if len(header) < 2 or header[0] != "parameter":
            continue
        # Expected converted table: parameter | suggested_explanation | file
        sug_col = None
        if "suggested_explanation" in header:
            sug_col = header.index("suggested_explanation")
        elif "suggested/corrected explanation" in header:
            sug_col = header.index("suggested/corrected explanation")
        elif len(header) >= 2:
            sug_col = 1
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if sug_col is None or len(cells) <= sug_col:
                continue
            p = cells[0].strip().upper()
            s = cells[sug_col].strip()
            if p and s and p != "-----------":
                out[p] = s
        if out:
            return out
    return out


def is_heading_paragraph(p: Paragraph) -> bool:
    st = (p.style.name or "").lower() if p.style else ""
    return st.startswith("heading")


def process_docx(path: Path, explain_map: dict[str, str]) -> tuple[bool, int]:
    doc = Document(str(path))
    paras = doc.paragraphs
    changed = False
    inserts = 0

    in_guidelines = False
    i = 0
    while i < len(paras):
        p = paras[i]
        txt = (p.text or "").strip()
        low = txt.lower()

        if low == "parameter configuration guidelines":
            in_guidelines = True
            i += 1
            continue

        if in_guidelines and low in STOP_TITLES:
            in_guidelines = False
            i += 1
            continue

        if not in_guidelines:
            i += 1
            continue

        m = RE_PARAM_HEAD.match(txt)
        if not m:
            i += 1
            continue

        param = m.group(1).upper()
        new_expl = explain_map.get(param)
        if not new_expl:
            i += 1
            continue

        # Capture explanation block: after heading until options / next param / section stop.
        j = i + 1
        explanation_idxs: list[int] = []
        found_same_new_text = False
        while j < len(paras):
            pj = paras[j]
            t = (pj.text or "").strip()
            l = t.lower()
            if not t:
                j += 1
                continue
            if l in STOP_TITLES:
                break
            if RE_PARAM_HEAD.match(t):
                break
            if RE_OPTIONS.match(t):
                break
            if t == new_expl:
                found_same_new_text = True
            explanation_idxs.append(j)
            j += 1

        if found_same_new_text:
            i = j
            continue

        if explanation_idxs:
            # Highlight existing explanation block in yellow.
            for idx in explanation_idxs:
                para = paras[idx]
                if not para.runs:
                    para.add_run("")
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            # Insert new explanation below existing explanation block.
            anchor = paras[explanation_idxs[-1]]
            insert_paragraph_after(anchor, new_expl, style_name=(anchor.style.name if anchor.style else None))
            changed = True
            inserts += 1
            # refresh list after insertion
            paras = doc.paragraphs
            i = j + 1
        else:
            # No explanation paragraph found; insert directly below heading.
            anchor = paras[i]
            insert_paragraph_after(anchor, new_expl, style_name=(anchor.style.name if anchor.style else None))
            changed = True
            inserts += 1
            paras = doc.paragraphs
            i += 2

    if changed:
        OUT_DIR.mkdir(parents=True, exist_ok=True)
        out_path = OUT_DIR / path.name
        doc.save(str(out_path))
    return changed, inserts


def main() -> None:
    dmap = load_dict_map(DICT_XLSX)
    gmap = load_ge3_map(GE3_DOCX)
    # Prefer GE3 curated explanations when param overlaps.
    merged = dict(dmap)
    merged.update(gmap)

    files = sorted(PART1_DIR.glob("*.docx"))
    changed_files = 0
    total_inserts = 0
    for fp in files:
        changed, inserts = process_docx(fp, merged)
        if changed:
            changed_files += 1
            total_inserts += inserts

    print(f"Part 1 files scanned: {len(files)}")
    print(f"Merged explanation parameters: {len(merged)}")
    print(f"Files updated: {changed_files}")
    print(f"Inserted new explanations: {total_inserts}")
    print(f"Output folder: {OUT_DIR}")


if __name__ == "__main__":
    main()

