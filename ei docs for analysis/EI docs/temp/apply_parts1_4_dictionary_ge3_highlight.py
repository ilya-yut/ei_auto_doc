"""
For each DOCX under EI docs/Part 1..Part 4:

- Use only params_dictionary.xlsx (sheet dictionary).
  Explanation text prefers column C (SAP canonical explanation), else column B.

- In "Parameter Configuration Guidelines", for parameters that exist in the dictionary:
  - Highlight existing explanation paragraphs in yellow
  - Insert the dictionary explanation below the highlighted block (one or more paragraphs
    if the cell text contains blank-line separators)

Outputs mirror folder layout under temp/Parts1_4_enriched/Part N/...

Each run deletes the output folder first so re-runs do not stack duplicate paragraphs.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook


ROOT = Path(r"c:\vibe code dev\ei_auto_doc")
EI_DOCS = ROOT / "ei docs for analysis" / "EI docs"
TEMP_DIR = EI_DOCS / "temp"
OUT_BASE = TEMP_DIR / "Parts1_4_enriched"

DICT_XLSX = TEMP_DIR / "params_dictionary.xlsx"

RE_PARAM_HEAD = re.compile(r"^([A-Z][A-Z0-9_]*)\s*\(.*\)\s*:?\s*$")
RE_OPTIONS = re.compile(r"^(?:\*?\*?)?[A-Z][A-Z0-9_]*\s*Options:\s*(?:\*?\*?)?$", re.I)
STOP_TITLES = {
    "parameter relationship",
    "parameter relationships",
    "default values",
    "practical example of parameter configuration",
    "practical configuration examples",
    "ei function structure",
    "abap code",
}


def insert_paragraph_after(paragraph: Paragraph, text: str, style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def load_dict_map(path: Path) -> dict[str, str]:
    wb = load_workbook(path, read_only=True)
    ws = wb["dictionary"] if "dictionary" in wb.sheetnames else wb[wb.sheetnames[0]]
    out: dict[str, str] = {}
    for r in range(2, ws.max_row + 1):
        p = ws.cell(r, 1).value
        if not p:
            continue
        param = str(p).strip().upper()
        c3 = ws.cell(r, 3).value
        c2 = ws.cell(r, 2).value
        canon = str(c3).strip() if c3 else ""
        sugg = str(c2).strip() if c2 else ""
        expl = canon or sugg
        if param and expl:
            out[param] = expl
    return out


def split_paragraphs(text: str) -> list[str]:
    parts = [x.strip() for x in text.split("\n\n") if x.strip()]
    return parts if parts else [text.strip()]


def process_docx(src: Path, out_dir: Path, explain_map: dict[str, str]) -> tuple[bool, int]:
    doc = Document(str(src))
    paras = doc.paragraphs
    changed = False
    inserts = 0

    in_guidelines = False
    i = 0
    while i < len(paras):
        p = paras[i]
        txt = (p.text or "").strip()
        low = txt.lower()

        if low == "parameter configuration guidelines":
            in_guidelines = True
            i += 1
            continue

        if in_guidelines and low in STOP_TITLES:
            in_guidelines = False
            i += 1
            continue

        if not in_guidelines:
            i += 1
            continue

        m = RE_PARAM_HEAD.match(txt)
        if not m:
            i += 1
            continue

        param = m.group(1).upper()
        new_expl = explain_map.get(param)
        if not new_expl:
            i += 1
            continue

        j = i + 1
        explanation_idxs: list[int] = []
        found_same_new_text = False
        pieces_check = split_paragraphs(new_expl)
        while j < len(paras):
            pj = paras[j]
            t = (pj.text or "").strip()
            l = t.lower()
            if not t:
                j += 1
                continue
            if l in STOP_TITLES:
                break
            if RE_PARAM_HEAD.match(t):
                break
            if RE_OPTIONS.match(t):
                break
            if t == new_expl.strip() or t in pieces_check:
                found_same_new_text = True
            explanation_idxs.append(j)
            j += 1

        if found_same_new_text:
            i = j
            continue

        pieces = split_paragraphs(new_expl)
        if explanation_idxs:
            for idx in explanation_idxs:
                para = paras[idx]
                if not para.runs:
                    para.add_run("")
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            anchor = paras[explanation_idxs[-1]]
            style_name = anchor.style.name if anchor.style else None
            cur = anchor
            for piece in pieces:
                cur = insert_paragraph_after(cur, piece, style_name=style_name)
            changed = True
            inserts += len(pieces)
            paras = doc.paragraphs
            i = j + 1
        else:
            anchor = paras[i]
            style_name = anchor.style.name if anchor.style else None
            cur = anchor
            for piece in pieces:
                cur = insert_paragraph_after(cur, piece, style_name=style_name)
            changed = True
            inserts += len(pieces)
            paras = doc.paragraphs
            i += 2

    if changed:
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / src.name
        doc.save(str(out_path))
    return changed, inserts


def main() -> None:
    if OUT_BASE.exists():
        shutil.rmtree(OUT_BASE)

    explain_map = load_dict_map(DICT_XLSX)

    parts = ["Part 1", "Part 2", "Part 3", "Part 4"]
    total_files = 0
    changed_files = 0
    total_inserts = 0

    for part in parts:
        src_dir = EI_DOCS / part
        if not src_dir.exists():
            print(f"Skip missing folder: {src_dir}")
            continue
        out_dir = OUT_BASE / part
        for fp in sorted(src_dir.glob("*.docx")):
            total_files += 1
            changed, ins = process_docx(fp, out_dir, explain_map)
            if changed:
                changed_files += 1
                total_inserts += ins

    print(f"Dictionary only: {DICT_XLSX}")
    print(f"Dictionary parameters with explanation: {len(explain_map)}")
    print(f"DOCX scanned (Part 1-4): {total_files}")
    print(f"Files written (modified): {changed_files}")
    print(f"Paragraph inserts: {total_inserts}")
    print(f"Output root: {OUT_BASE}")


if __name__ == "__main__":
    main()
