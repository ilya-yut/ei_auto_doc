from __future__ import annotations

import re
import sys
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font


ROOT = Path(r"c:\vibe code dev\ei_auto_doc")
PART1_DOCX_DIR = ROOT / "ei docs for analysis" / "EI docs" / "Part 1"
TEMP_DIR = ROOT / "ei docs for analysis" / "EI docs" / "temp"

SOURCE_DICTIONARY_XLSX = TEMP_DIR / "params_dictionary.xlsx"
SOURCE_GE3_DOCX = (
    ROOT
    / "ei docs for analysis"
    / "Part 1 conv"
    / "SHARED_PARAMETER_EXPLANATIONS_GE3_NO_BACKDAYS_USER_FLD.docx"
)
OUT_XLSX = TEMP_DIR / "params_dictionary_part1_missing.xlsx"

RE_NUM = re.compile(r"^\d+$")
RE_PARAM_TOKEN = re.compile(r"^[A-Z][A-Z0-9_]{1,}$")


def iter_block_items(parent: DocumentType):
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_params_from_part1_docx(path: Path) -> tuple[set[str], dict[str, dict[str, str]]]:
    doc = Document(str(path))
    saw_param_ref_heading = False
    out: set[str] = set()
    meta: dict[str, dict[str, str]] = {}

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            t = block.text.strip().lower()
            if "parameters reference table" in t:
                saw_param_ref_heading = True
            continue
        if not isinstance(block, Table):
            continue
        if not saw_param_ref_heading:
            continue

        rows = [[c.text.strip() for c in row.cells] for row in block.rows]
        if not rows:
            continue
        low_header = [x.lower() for x in rows[0]]
        if "field" not in low_header:
            continue
        field_col = low_header.index("field")
        desc_col = low_header.index("description") if "description" in low_header else None
        type_col = low_header.index("type") if "type" in low_header else None
        de_col = low_header.index("data element") if "data element" in low_header else None
        dom_col = low_header.index("domain") if "domain" in low_header else None
        num_col = None
        if "#" in rows[0]:
            num_col = rows[0].index("#")
        elif "no." in low_header:
            num_col = low_header.index("no.")
        elif "no" in low_header:
            num_col = low_header.index("no")

        for r in rows[1:]:
            if field_col >= len(r):
                continue
            field = r[field_col].strip().upper()
            if not field or field in {"FIELD", "---"}:
                continue
            if not RE_PARAM_TOKEN.match(field):
                continue
            if num_col is not None and num_col < len(r):
                n = r[num_col].strip()
                if n and not RE_NUM.match(n):
                    break
            out.add(field)
            meta[field] = {
                "description": (r[desc_col].strip() if desc_col is not None and desc_col < len(r) else ""),
                "type": (r[type_col].strip() if type_col is not None and type_col < len(r) else ""),
                "data_element": (r[de_col].strip() if de_col is not None and de_col < len(r) else ""),
                "domain": (r[dom_col].strip() if dom_col is not None and dom_col < len(r) else ""),
            }
        # first parameters table after heading is enough
        if out:
            return out, meta
    return out, meta


def read_params_from_dictionary_xlsx(path: Path) -> set[str]:
    wb = load_workbook(path, read_only=True)
    ws = wb[wb.sheetnames[0]]
    params: set[str] = set()
    # assume first row header: parameter
    for r in range(2, ws.max_row + 1):
        v = ws.cell(r, 1).value
        if v is None:
            continue
        s = str(v).strip().upper()
        if s:
            params.add(s)
    return params


def read_params_from_ge3_docx(path: Path) -> set[str]:
    doc = Document(str(path))
    params: set[str] = set()
    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        if not rows:
            continue
        header = [x.lower() for x in rows[0]]
        if len(header) >= 1 and header[0] == "parameter":
            for r in rows[1:]:
                if not r:
                    continue
                p = r[0].strip().upper()
                if p and p != "-----------":
                    params.add(p)
            if params:
                return params
    return params


def build_explainer():
    sys.path.insert(0, str(TEMP_DIR))
    from build_params_dictionary_xlsx import OVERRIDES  # type: ignore

    return OVERRIDES


def _pick_best(counter: Counter[str]) -> str:
    for k, _ in counter.most_common():
        s = k.strip()
        if not s:
            continue
        low = s.lower()
        if "not used" in low or "unused" in low:
            continue
        return s
    return counter.most_common(1)[0][0].strip() if counter else ""


def _pattern_explanation(param: str) -> str:
    p = param.upper()
    m = re.match(r"BP([123])_(CODE|FUNCT|NAME)$", p)
    if m:
        slot, kind = m.groups()
        if kind == "CODE":
            return f"{p} identifies business partner slot {slot} code used in multi-partner comparisons."
        if kind == "FUNCT":
            return f"{p} defines business partner slot {slot} role/function for partner responsibility analysis."
        return f"{p} carries business partner slot {slot} name for readable partner-level reporting."
    if p.endswith("_DESC"):
        base = p[: -5]
        return f"{p} provides description text for {base}, used for business-readable reporting."
    if p.endswith("_OLD"):
        base = p[: -4]
        return f"{p} stores the previous value of {base} for before/after change analysis."
    if p.endswith("_NEW"):
        base = p[: -4]
        return f"{p} stores the new value of {base} for after-change impact analysis."
    if p.endswith("DATE") or p.endswith("DAT"):
        return f"{p} is a date selector used to scope records to a relevant monitoring period."
    if p.endswith("TIME") or p.endswith("TIM"):
        return f"{p} is a time selector used to refine event windows within the selected dates."
    if p.endswith("ICON"):
        return f"{p} is a status icon field used to present state/severity in output."
    if re.search(r"\bCNT\b", p) or p.endswith("_CNT"):
        return f"{p} is a count metric used for threshold-based exception evaluation."
    if p in {"KUNNR", "LIFNR", "MATNR", "VBELN", "EBELN", "BUKRS", "WERKS", "VKORG"}:
        # Should usually be covered in OVERRIDES, but keep explicit business-safe fallback.
        return f"{p} is a core SAP key used to scope records to the relevant business object set."
    return f"{p} is a technical SAP field used in this monitor context; align interpretation with DDIC object semantics."


def build_explanation_for_param(param: str, metas: list[dict[str, str]], overrides: dict[str, str]) -> str:
    p = param.upper()
    if p in overrides:
        return overrides[p]

    desc_c: Counter[str] = Counter()
    de_c: Counter[str] = Counter()
    dom_c: Counter[str] = Counter()
    typ_c: Counter[str] = Counter()
    for m in metas:
        d = (m.get("description") or "").strip()
        if d:
            desc_c[d] += 1
        de = (m.get("data_element") or "").strip()
        if de:
            de_c[de] += 1
        dom = (m.get("domain") or "").strip()
        if dom:
            dom_c[dom] += 1
        typ = (m.get("type") or "").strip()
        if typ:
            typ_c[typ] += 1

    desc = _pick_best(desc_c)
    de = _pick_best(de_c)
    dom = _pick_best(dom_c)
    typ = _pick_best(typ_c)

    if desc:
        dlow = desc.lower().rstrip(".")
        if "count" in dlow:
            base = f"{p} defines threshold/range control for {dlow}"
        elif "date" in dlow:
            base = f"{p} scopes records by {dlow}"
        elif "time" in dlow:
            base = f"{p} refines event selection by {dlow}"
        elif "status" in dlow or "state" in dlow:
            base = f"{p} filters records by {dlow}"
        elif "user" in dlow:
            base = f"{p} restricts records by {dlow}"
        elif "currency" in dlow:
            base = f"{p} controls currency-context filtering by {dlow}"
        elif "amount" in dlow or "value" in dlow:
            base = f"{p} is used for value-based filtering ({dlow})"
        else:
            base = f"{p} scopes records by {dlow}"

        ddic_bits = [x for x in [de, dom] if x]
        if ddic_bits:
            base += f" (DDIC: {' / '.join(ddic_bits)})"
        return base + "."

    # No usable table description -> robust pattern fallback.
    pat = _pattern_explanation(p)
    if de or dom or typ:
        dd = " / ".join([x for x in [de, dom, typ] if x])
        return f"{pat} (DDIC hint: {dd})."
    return pat


def main() -> None:
    part1_files = sorted(PART1_DOCX_DIR.glob("*.docx"))
    if not part1_files:
        raise FileNotFoundError(f"No .docx files found in {PART1_DOCX_DIR}")

    part1_params: dict[str, set[str]] = defaultdict(set)
    meta_by_param: dict[str, list[dict[str, str]]] = defaultdict(list)
    skipped: list[str] = []
    for fp in part1_files:
        params, meta = extract_params_from_part1_docx(fp)
        if not params:
            skipped.append(fp.name)
            continue
        for p in params:
            part1_params[p].add(fp.name)
            if p in meta:
                meta_by_param[p].append(meta[p])

    dict_params = read_params_from_dictionary_xlsx(SOURCE_DICTIONARY_XLSX)
    ge3_params = read_params_from_ge3_docx(SOURCE_GE3_DOCX)
    existing = dict_params | ge3_params

    missing = sorted([p for p in part1_params.keys() if p not in existing])
    overrides = build_explainer()

    wb = Workbook()
    ws = wb.active
    ws.title = "missing_dictionary"

    ws["A1"] = "Summary"
    ws["A1"].font = Font(bold=True)
    ws["A2"] = "Part 1 DOCX scanned"
    ws["B2"] = len(part1_files)
    ws["A3"] = "Part 1 distinct parameters"
    ws["B3"] = len(part1_params)
    ws["A4"] = "Existing params (dictionary ∪ ge3 docx)"
    ws["B4"] = len(existing)
    ws["A5"] = "Missing params (to add)"
    ws["B5"] = len(missing)
    ws["A6"] = "Skipped Part 1 files (no parsed params table)"
    ws["B6"] = len(skipped)

    header_row = 8
    ws.cell(header_row, 1, "parameter").font = Font(bold=True)
    ws.cell(header_row, 2, "suggested/corrected explanation").font = Font(bold=True)
    ws.cell(header_row, 3, "used_in_part1_files_count").font = Font(bold=True)

    r = header_row + 1
    for p in missing:
        ws.cell(r, 1, p)
        ws.cell(r, 2, build_explanation_for_param(p, meta_by_param.get(p, []), overrides))
        ws.cell(r, 3, len(part1_params[p]))
        r += 1

    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 130
    ws.column_dimensions["C"].width = 24

    wb.save(OUT_XLSX)

    print(f"Part 1 DOCX scanned: {len(part1_files)}")
    print(f"Part 1 distinct params: {len(part1_params)}")
    print(f"Existing params (dictionary + ge3): {len(existing)}")
    print(f"Missing params written: {len(missing)}")
    print(f"Skipped files: {len(skipped)}")
    if skipped:
        print("Skipped file examples:", ", ".join(skipped[:5]))
    print(f"Wrote {OUT_XLSX}")


if __name__ == "__main__":
    main()

