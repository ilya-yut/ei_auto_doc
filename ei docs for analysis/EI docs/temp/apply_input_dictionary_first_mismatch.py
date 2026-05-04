"""Align Parameter Configuration Guidelines with input/params_dictionary.xlsx (in place).

- For each `PARAM (...):` block under Parameter Configuration Guidelines whose name exists in the
  dictionary (columns parameter + suggested/corrected explanation):
  - If the block already contains the dictionary wording (verbatim paragraph match) -> skip.
  - If normalized joined explanation equals the dictionary text -> skip.
  - Else -> highlight existing explanation paragraphs in yellow, insert dictionary text below.

Default: process every mismatching parameter (reload document after each save so indices stay valid).
Pass `--first-only` to stop after the first fix (legacy behavior).
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

ROOT = Path(r"c:\vibe code dev\ei_auto_doc")

DICT = ROOT / "input" / "params_dictionary.xlsx"
DEFAULT_DOC = (
    ROOT
    / "ei docs for analysis"
    / "EI docs"
    / "temp"
    / "Parts1_4_enriched"
    / "Part 2"
    / "SW_01_20_USER_STATE - User Actions Control.docx"
)

RE_PARAM_HEAD = re.compile(r"^([A-Z][A-Z0-9_]*)\s*\(.*\)\s*:?\s*$")
RE_OPTIONS = re.compile(r"^(?:\*?\*?)?[A-Z][A-Z0-9_]*\s*Options:\s*(?:\*?\*?)?$", re.I)
STOP = {
    "parameter relationship",
    "parameter relationships",
    "default values",
    "practical example of parameter configuration",
    "practical configuration examples",
    "ei function structure",
    "abap code",
}


def load_input_dict(path: Path) -> dict[str, str]:
    wb = load_workbook(path, read_only=True)
    ws = wb["dictionary"] if "dictionary" in wb.sheetnames else wb[wb.sheetnames[0]]
    out: dict[str, str] = {}
    for r in range(2, ws.max_row + 1):
        p = ws.cell(r, 1).value
        if not p:
            continue
        k = str(p).strip().upper()
        v = ws.cell(r, 2).value
        t = str(v).strip() if v else ""
        if k and t:
            out[k] = (
                t.replace("\ufffd", "'")
                .replace("\u2019", "'")
                .replace("\u2013", "-")
                .replace("\u2014", "-")
            )
    wb.close()
    return out


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def split_paragraphs(text: str) -> list[str]:
    parts = [x.strip() for x in text.split("\n\n") if x.strip()]
    return parts if parts else [text.strip()]


def insert_paragraph_after(
    paragraph: Paragraph, text: str, style_name: str | None = None
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def try_one_fix(doc: Document, explain_map: dict[str, str]) -> tuple[bool, str]:
    """Apply at most one mismatch fix; return (changed, message)."""
    paras = doc.paragraphs
    in_g = False
    i = 0
    while i < len(paras):
        t = (paras[i].text or "").strip()
        low = t.lower()
        if low == "parameter configuration guidelines":
            in_g = True
            i += 1
            continue
        if in_g and low in STOP:
            break
        if not in_g:
            i += 1
            continue
        m = RE_PARAM_HEAD.match(t)
        if not m:
            i += 1
            continue
        param = m.group(1).upper()
        ref = explain_map.get(param)
        if not ref:
            i += 1
            continue
        j = i + 1
        expl_idxs: list[int] = []
        found_dict_paragraph = False
        pieces_check = split_paragraphs(ref)
        while j < len(paras):
            t2 = (paras[j].text or "").strip()
            l2 = t2.lower()
            if not t2:
                j += 1
                continue
            if l2 in STOP:
                break
            if RE_PARAM_HEAD.match(t2):
                break
            if RE_OPTIONS.match(t2):
                break
            if t2 == ref.strip() or t2 in pieces_check:
                found_dict_paragraph = True
            expl_idxs.append(j)
            j += 1

        if found_dict_paragraph:
            i = j
            continue

        doc_text = " ".join((paras[k].text or "").strip() for k in expl_idxs)
        if norm(doc_text) == norm(ref):
            i = j
            continue

        pieces = split_paragraphs(ref)
        if expl_idxs:
            for idx in expl_idxs:
                para = paras[idx]
                if not para.runs:
                    para.add_run("")
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            anchor = paras[expl_idxs[-1]]
        else:
            anchor = paras[i]
        style_name = anchor.style.name if anchor.style else None
        cur = anchor
        for piece in pieces:
            cur = insert_paragraph_after(cur, piece, style_name=style_name)
        msg = f"UPDATED param {param} yellow_paras={len(expl_idxs)} inserted={len(pieces)}"
        return True, msg

    return False, ""


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", nargs="?", default=None, help="Path to DOCX (default: USER_STATE Part 2)")
    ap.add_argument(
        "--first-only",
        action="store_true",
        help="Stop after the first applied fix (single save).",
    )
    args = ap.parse_args()

    doc_path = Path(args.docx) if args.docx else DEFAULT_DOC
    first_only = bool(args.first_only)

    if not DICT.is_file():
        print("Missing dictionary:", DICT)
        return 1
    if not doc_path.is_file():
        print("Missing docx:", doc_path)
        return 1

    explain_map = load_input_dict(DICT)
    total_fixes = 0
    while True:
        doc = Document(str(doc_path))
        changed, msg = try_one_fix(doc, explain_map)
        if not changed:
            break
        doc.save(str(doc_path))
        total_fixes += 1
        print(msg)
        if first_only:
            break

    if total_fixes == 0:
        print("No further mismatching parameters (guidelines + dictionary).")
    else:
        print(f"Done. Total fixes applied: {total_fixes}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
