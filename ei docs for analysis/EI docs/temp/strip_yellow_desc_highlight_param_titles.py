"""
Under Parameter Configuration Guidelines:

1. Parameter must be in an allow-set: table[0] Field column (default), or `--allowlist kernel`,
   or `--allowlist ord-price-m` (SD order / pricing monitor list).
2. For blocks whose explanation paragraphs contain at least one run with yellow highlight:
   - Remove only runs that use yellow highlight (leave other runs untouched).
   - Delete explanation paragraphs that are empty after stripping.
3. Set yellow highlight on the parameter title paragraph (`FIELD (...):`).
4. Remove any remaining empty paragraphs inside the guidelines section.

In-place save. Use `--directory FOLDER` to process every `.docx` in that folder (skips `~$` lock files).
Use `--mtime-today-before HH:MM` with `--directory` to only process files whose last modification time
is today's local date and strictly before that clock time.
Use `--mtime-today-at HH:MM` (not together with --mtime-today-before) to only process files modified
during that clock minute on today's local date.

Use `--merge-from-xlsx PATH.xlsx` together with `--directory FOLDER` to, for each DOCX, find parameter
blocks under "Parameter Configuration Guidelines" whose name exists in the Excel dictionary sheet
`dictionary` (columns: parameter, suggested/corrected explanation), highlight all runs in the existing
explanation paragraphs yellow, then append the dictionary explanation below (plain text, new
paragraphs). In-place save. Requires openpyxl.

Use ``--normalize-practical-examples`` with ``--directory`` to rewrite each "Practical Configuration Examples"
use case as: bold title, then Purpose (all non-bold), then one paragraph per ``KEY = value`` line, then
one blank paragraph. Only reorders/splits existing text (no new parameters).
"""

from __future__ import annotations

import argparse
import re
import sys
from datetime import date, datetime, time
from pathlib import Path
from typing import cast

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(r"c:\vibe code dev\ei_auto_doc")
DEFAULT_DOC = (
    ROOT
    / "ei docs for analysis"
    / "EI docs"
    / "temp"
    / "Parts1_4_enriched"
    / "Part 2"
    / "SW_01_20_USER_STATE - User Actions Control.docx"
)

# User-supplied comparison list (Kernel / user-master style monitors). INT_PWD_ICON is a common
# alternate spelling for INIT_PWD_ICON in docs — both match when this list is active.
EXPLICIT_PARAM_ALLOWLIST: frozenset[str] = frozenset(
    {
        "ACCNT",
        "ANAME",
        "BCDA1",
        "BNAME",
        "CLASS",
        "CODV1",
        "CODVN",
        "DATE_REF_FLD",
        "DURATION",
        "DURATION_UNIT",
        "ERDAT",
        "GLTGB",
        "GLTGV",
        "INT_PWD_ICON",
        "INIT_PWD_ICON",
        "LANGU",
        "LOCK_ICON",
        "MODDATE",
        "MODIFIER",
        "MODTIME",
        "NO_DATE_RESTRICTION",
        "PWDLGNDATE",
        "STATE_COLOR",
        "STATE_DESC",
        "STATE_ICON",
        "TRDAT",
        "TZONE",
        "UFLAG",
        "USTYP",
        "VALID_USERS_ONLY",
        "VERSN",
    }
)

# SD / MM / FI style monitors (e.g. Part 3/4) — Field-style names + INT_PWD_ICON alias for INIT_PWD_ICON.
ALLOWLIST_ORD_PRICE_M: frozenset[str] = frozenset(
    {
        "ABCIN",
        "ACCNT",
        "ACTFLG",
        "ACTION",
        "ACTION_DESC",
        "ACT_CHNGNO",
        "AEDAT",
        "AENAM",
        "ANAME",
        "ARKTX",
        "AS4DATE",
        "AS4TIME",
        "AUART",
        "AUDAT",
        "AUFNR",
        "AUSBK",
        "AWORG",
        "AWREF",
        "AWSYS",
        "AWTYP",
        "BADAT",
        "BANFN",
        "BANKL",
        "BANKN",
        "BANKS",
        "BCDA1",
        "BEDAT",
        "BELNR",
        "BKLAS",
        "BKTXT",
        "BLART",
        "BLDAT",
        "BNAME",
        "BNFPO",
        "BPMNG",
        "BPRME",
        "BPUMN",
        "BPUMZ",
        "BRGEW",
        "BSAKZ",
        "BSART",
        "BSART_DESC",
        "BSTAR",
        "BSTAT",
        "BSTNK",
        "BSTYP",
        "BSTYP_DESC",
        "BUDAT",
        "BUKRS",
        "BUTXT",
        "BUZEI",
        "BWKEY",
        "BWTTY",
        "BZIRK",
        "CHANGEIND_DESC",
        "CHANGENR",
        "CHANGE_IND",
        "CHANGE_IND_DESC",
        "CHARG",
        "CHNGIND",
        "CHNGIND_DESC",
        "CLASS",
        "CMGST",
        "CODV1",
        "CODVN",
        "COMP_CODE",
        "COMP_CODE_DESC",
        "COMP_OPERATOR",
        "COMSYS",
        "COND",
        "COSTA",
        "CPUDT",
        "CUKY_NEW",
        "CUKY_OLD",
        "CUST_DESC",
        "DATE_REF_FLD",
        "DEPARTMENT",
        "DEST",
        "DEVCLASS",
        "DIWZL",
        "DMBTR",
        "DSTAT",
        "DURATION",
        "DURATION_H",
        "DURATION_M",
        "DURATION_UNIT",
        "EBELN",
        "EBELP",
        "EDATU",
        "EERNO",
        "EINDT",
        "EKGRP",
        "EKGRP_DESC",
        "EKNAM",
        "EKORG",
        "EKORG_DESC",
        "EKOTX",
        "ELIKZ",
        "ERDAT",
        "EREKZ",
        "ERFME",
        "ERFMG",
        "ERNAM",
        "ERZET",
        "ESTKZ",
        "ETENR",
        "EXVKW",
        "FABKL",
        "FACDATE",
        "FAKSK",
        "FDGRV",
        "FKART",
        "FKDAT",
        "FKSTA",
        "FKSTK",
        "FKSTO",
        "FKTYP",
        "FMSTK",
        "FNAME",
        "FRGC",
        "FRGGR",
        "FRGKE",
        "FRGKZ",
        "FRGRL",
        "FRGST",
        "FRGSX",
        "FRGZU",
        "FROMNUMBER",
        "GARG",
        "GCLIENT",
        "GEWEI",
        "GLTGB",
        "GLTGV",
        "GMODE",
        "GRUND",
        "GSBER",
        "GTCODE",
        "GTHOST",
        "GUIVERSION",
        "HKONT",
        "HOST",
        "HOSTADR",
        "HWAER",
        "IBLNR",
        "IMPFLAG",
        "IMPFLG",
        "IMPSING",
        "INIT_PWD_ICON",
        "INSTANCENAME",
        "INT_PWD_ICON",
        "IP_ADDRESS",
        "JOBID",
        "KDATB",
        "KDATE",
        "KDAUF",
        "KDGRP",
        "KDPOS",
        "KLMENG",
        "KNTTP",
        "KNUMV",
        "KOBIS",
        "KOKRS",
        "KONZS",
        "KOSTL",
        "KOVON",
        "KSCHL",
        "KTOKK",
        "KTOPL",
        "KUNAG",
        "KUNNR",
        "KUNRG",
        "KURSF",
        "KWERT",
        "KWMENG",
        "KZKRS",
        "KZWRS",
        "LAEDA",
        "LAND1",
        "LANG",
        "LANGU",
        "LASTCHNAME",
        "LBKUM",
        "LDDAT",
        "LFART",
        "LFBNR",
        "LFDAT",
        "LFGSA",
        "LFIMG",
        "LFM1_LOEVM",
        "LFM1_SPERM",
        "LFPOS",
        "LFSTA",
        "LGMNG",
        "LGORT",
        "LIFNR",
        "LOCK_ICON",
        "LOEKZ",
        "LOGSYS",
        "LSTAT",
        "LVSTA",
        "MANAGE_IN_UTC",
        "MANDT",
        "MATKL",
        "MATNR",
        "MAT_DESC",
        "MBLNR",
        "MEINS",
        "MEMSUM",
        "MENGE",
        "MJAHR",
        "MODDA",
        "MODDATE",
        "MODIFIER",
        "MODTI",
        "MODTIME",
        "MONAT",
        "MPROK",
        "MPROK_DESC",
        "MSCDATE",
        "MSCTIME",
        "MSGID",
        "MSGNO",
        "MSTYP",
        "MTART",
        "NETWR",
        "NEW_VAL",
        "NODATAFLG",
        "NO_DATE_RESTRICTION",
        "NTGEW",
        "OBJECTCLAS",
        "OBJECT_DESC",
        "OBJNR",
        "OLD_VAL",
        "PARGB",
        "PARVW",
        "PAYER_DESC",
        "PEINH",
        "PKSTA",
        "PKSTK",
        "PLANCHNGNR",
        "PLANT_DESC",
        "PLIFZ",
        "PODAT",
        "POSNR",
        "PRCTR",
        "PREFLG",
        "PRIVSUM",
        "PROCSTAT",
        "PROCSTAT_DESC",
        "PROF_ASS",
        "PROF_ASS_T",
        "PROF_DEL",
        "PROGNAME",
        "PROTOCOL",
        "PSTYP",
        "PSTYV",
        "PS_PSP_PNR",
        "PWDLGNDATE",
        "PWDLOCKDATE",
        "RELIK",
        "RESWK",
        "RESWK_DESC",
        "RFBSK",
        "RFCDEST",
        "RFC_TYPE",
        "RLWRT",
        "ROLL",
        "ROUTE",
        "SAKTO",
        "SALK3",
        "SDLSTRTDT",
        "SDLSTRTTM",
        "SDLUNAME",
        "SGTXT",
        "SHIPTO_DESC",
        "SHKZG",
        "SLGMAND",
        "SLGMODE",
        "SLGPROC",
        "SLGREPNA",
        "SLGTC",
        "SLGTIME",
        "SLGTYPE",
        "SLGUSER",
        "SOBKZ",
        "SOLDTO_DESC",
        "SPART",
        "SPERR",
        "SPRAS",
        "STAT",
        "STATE_COLOR",
        "STATE_DESC",
        "STATE_ICON",
        "STATU",
        "STATUS",
        "STATUS_DESC",
        "STATU_DESC",
        "STKZN",
        "STPRS",
        "SW_DEST",
        "SYSNAM",
        "TAB",
        "TABKEY",
        "TABNAME",
        "TAB_DESC",
        "TARCLI",
        "TARSYSTEM",
        "TCODE",
        "TERM",
        "TEXT_CASE",
        "TID",
        "TIME_DIFF",
        "TONUMBER",
        "TOTAL_MEM_MB",
        "TOYEAR",
        "TRDAT",
        "TRFUNCTION",
        "TRKORR",
        "TRSTA",
        "TRSTATUS",
        "TXZ01",
        "TYPE",
        "TZONE",
        "UDATE",
        "UEBTK",
        "UEBTO",
        "UFLAG",
        "UMSAV",
        "UNAME",
        "UNIT_NEW",
        "UNIT_OLD",
        "UPD_DATE",
        "UPD_TIME",
        "USER",
        "USERID",
        "USERNAME",
        "USNAD",
        "USNAM",
        "USTYP",
        "USTYP_DESC",
        "UTIME",
        "UVPAS",
        "UVPIS",
        "UZNAZ",
        "VALID_USERS_ONLY",
        "VALUE_NEW",
        "VALUE_OLD",
        "VBELN",
        "VBELP",
        "VBTYP",
        "VBUND",
        "VDATU",
        "VENDOR_DESC",
        "VERPR",
        "VERSN",
        "VESTK",
        "VGABE",
        "VGART",
        "VGBEL",
        "VGPOS",
        "VGTYP",
        "VKBUR",
        "VKGRP",
        "VKMZL",
        "VKNZL",
        "VKORG",
        "VKWRA",
        "VKWRT",
        "VLSTK",
        "VOLEH",
        "VOLUM",
        "VPRSV",
        "VRGNG",
        "VRKME",
        "VSTEL",
        "VTEXT",
        "VTWEG",
        "WADAT",
        "WADAT_IST",
        "WAERK",
        "WAERS",
        "WAERS_FR",
        "WAS_PLANND",
        "WAVWR",
        "WBSTA",
        "WBSTK",
        "WEPOS",
        "WERKS",
        "WERKS_DESC",
        "WGBEZ",
        "WORKING_DAYS",
        "WP_TYPE",
        "WRBTR",
        "WRTBM",
        "WRTZL",
        "XAMEI",
        "XBLNI",
        "XBLNR",
        "XCPDK",
        "XNULL",
        "XNZAE",
        "XZAEL",
        "ZEILE",
        "ZEIT",
        "ZTERM",
    }
)
ALLOWLIST_BY_NAME: dict[str, frozenset[str]] = {
    "kernel": EXPLICIT_PARAM_ALLOWLIST,
    "ord-price-m": ALLOWLIST_ORD_PRICE_M,
}

RE_PARAM_HEAD = re.compile(r"^([A-Z][A-Z0-9_]*)\s*\(.*\)\s*:?\s*$")
RE_OPTIONS = re.compile(r"^(?:\*?\*?)?[A-Z][A-Z0-9_]*\s*Options:\s*(?:\*?\*?)?$", re.I)
STOP = {
    "parameter relationship",
    "parameter relationships",
    "default values",
    "practical example of parameter configuration",
    "practical configuration examples",
    "ei function structure",
    "abap code",
}

# Headings that end the "Practical Configuration Examples" body (after the section title).
END_PRACTICAL_CONFIGURATION_EXAMPLES_SECTION = frozenset(
    {
        "default values",
        "parameter relationship",
        "parameter relationships",
        "practical example of parameter configuration",
        "ei function structure",
        "abap code",
    }
)

RE_USE_CASE_TITLE = re.compile(r"^Use Case\s+\d+\s*:\s*.+", re.I | re.DOTALL)


def reference_field_set(doc: Document) -> set[str]:
    if not doc.tables:
        return set()
    t0 = doc.tables[0]
    out: set[str] = set()
    for ri in range(1, len(t0.rows)):
        row = t0.rows[ri]
        if len(row.cells) < 2:
            continue
        f = (row.cells[1].text or "").strip().upper()
        if f:
            out.add(f)
    return out


def paragraph_has_yellow_run(p: Paragraph) -> bool:
    for r in p.runs:
        if r.font.highlight_color == WD_COLOR_INDEX.YELLOW:
            return True
    return False


def strip_yellow_runs(p: Paragraph) -> int:
    """Remove w:r elements whose highlight is yellow. Returns number of runs removed."""
    removed = 0
    for run in list(p.runs):
        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
            run._element.getparent().remove(run._element)
            removed += 1
    return removed


def delete_paragraph(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_blank_paragraphs_in_guidelines(doc: Document) -> int:
    """Delete empty paragraphs between guidelines heading and first STOP title. Returns count removed."""
    removed = 0
    while True:
        paras = doc.paragraphs
        start = end = None
        in_g = False
        for i, p in enumerate(paras):
            t = (p.text or "").strip()
            low = t.lower()
            if low == "parameter configuration guidelines":
                in_g = True
                start = i
                continue
            if in_g and low in STOP:
                end = i
                break
        if start is None or end is None:
            break
        deleted = False
        for i in range(start + 1, end):
            p = paras[i]
            if not (p.text or "").strip():
                delete_paragraph(p)
                removed += 1
                deleted = True
                break
        if not deleted:
            break
    return removed


def mtime_today_before(fp: Path, hh_mm: str | None) -> bool:
    """True if no filter, else file mtime is local today and strictly before HH:MM."""
    if not hh_mm:
        return True
    parts = hh_mm.strip().split(":")
    if len(parts) != 2:
        raise SystemExit("--mtime-today-before must be HH:MM (24h)")
    h, m = int(parts[0]), int(parts[1])
    boundary = datetime.combine(date.today(), time(h, m))
    mt = datetime.fromtimestamp(fp.stat().st_mtime)
    return mt.date() == date.today() and mt < boundary


def mtime_today_at(fp: Path, hh_mm: str | None) -> bool:
    """True if no filter, else local mtime is today and falls in the same clock minute as HH:MM."""
    if not hh_mm:
        return True
    parts = hh_mm.strip().split(":")
    if len(parts) != 2:
        raise SystemExit("--mtime-today-at must be HH:MM (24h)")
    h, m = int(parts[0]), int(parts[1])
    mt = datetime.fromtimestamp(fp.stat().st_mtime)
    return mt.date() == date.today() and mt.hour == h and mt.minute == m


def highlight_title_yellow(p: Paragraph) -> None:
    if not p.runs:
        p.add_run(p.text or "")
    for run in p.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def load_params_dictionary(xlsx_path: Path) -> dict[str, str]:
    try:
        import openpyxl
    except ImportError as e:
        raise SystemExit("openpyxl is required for --merge-from-xlsx") from e
    if not xlsx_path.is_file():
        raise SystemExit(f"Dictionary not found: {xlsx_path}")
    wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
    if "dictionary" not in wb.sheetnames:
        wb.close()
        raise SystemExit("Expected sheet named 'dictionary' in the workbook.")
    ws = wb["dictionary"]
    rows = ws.iter_rows(values_only=True)
    header = next(rows, None)
    if not header or len(header) < 2:
        wb.close()
        raise SystemExit("Dictionary sheet must have at least two columns.")
    out: dict[str, str] = {}
    for row in rows:
        if not row or row[0] is None:
            continue
        key = str(row[0]).strip().upper()
        if not key:
            continue
        val = "" if row[1] is None else str(row[1]).strip()
        out[key] = val
    wb.close()
    return out


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    else:
        new_para.add_run("")
    return new_para


def highlight_explanation_paragraphs_yellow(paragraphs: list[Paragraph]) -> None:
    for p in paragraphs:
        if not (p.text or "").strip():
            continue
        for run in p.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def collect_guideline_param_blocks(doc: Document) -> list[tuple[int, list[int], str]]:
    """Return list of (title_para_index, explanation_para_indices, PARAM_UPPER) within guidelines."""
    paras = doc.paragraphs
    in_g = False
    i = 0
    blocks: list[tuple[int, list[int], str]] = []
    while i < len(paras):
        t = (paras[i].text or "").strip()
        low = t.lower()
        if low == "parameter configuration guidelines":
            in_g = True
            i += 1
            continue
        if in_g and low in STOP:
            break
        if not in_g:
            i += 1
            continue
        m = RE_PARAM_HEAD.match(t)
        if not m:
            i += 1
            continue
        param = m.group(1).upper()
        title_i = i
        j = i + 1
        expl_idxs: list[int] = []
        while j < len(paras):
            t2 = (paras[j].text or "").strip()
            l2 = t2.lower()
            if not t2:
                j += 1
                continue
            if l2 in STOP:
                break
            if RE_PARAM_HEAD.match(t2):
                break
            if RE_OPTIONS.match(t2):
                break
            expl_idxs.append(j)
            j += 1
        blocks.append((title_i, expl_idxs, param))
        i = j
    return blocks


def is_param_assignment_line(s: str) -> bool:
    s = s.strip()
    if "=" not in s or s.lower().startswith("purpose:"):
        return False
    left, _, _right = s.partition("=")
    left = left.strip()
    if not left or not re.fullmatch(r"[A-Z0-9_]+", left, re.I):
        return False
    return True


def classify_practical_body_paragraph(p: Paragraph) -> str:
    text = (p.text or "").strip()
    if not text:
        return "empty"
    if text.lower().startswith("purpose:"):
        return "purpose"
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if lines and all(is_param_assignment_line(ln) for ln in lines):
        return "params"
    return "other"


def normalize_assignment_spacing(line: str) -> str:
    left, _, right = line.partition("=")
    if "=" not in line:
        return line.strip()
    return f"{left.strip()} = {right.strip()}"


def paragraph_body_index(doc: Document, p: Paragraph) -> int:
    for i, q in enumerate(doc.paragraphs):
        if q._element is p._element:
            return i
    return -1


def set_paragraph_bold_all_runs(p: Paragraph, bold: bool | None) -> None:
    if not p.runs and (p.text or "").strip():
        p.add_run(p.text or "")
    if not p.runs:
        return
    for r in p.runs:
        r.bold = bold


def collect_practical_configuration_example_blocks(doc: Document) -> list[dict]:
    blocks: list[dict] = []
    cur: dict | None = None
    in_sec = False
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        low = t.lower()
        if t == "Practical Configuration Examples":
            in_sec = True
            continue
        if in_sec and low in END_PRACTICAL_CONFIGURATION_EXAMPLES_SECTION:
            break
        if not in_sec:
            continue
        if RE_USE_CASE_TITLE.match(t):
            if cur is not None:
                blocks.append(cur)
            cur = {"title": p, "middle": []}
            continue
        if cur is not None:
            cur["middle"].append(p)
    if cur is not None:
        blocks.append(cur)
    return blocks


def normalize_practical_configuration_examples(doc_path: Path) -> int:
    """
    For each ``Use Case N:`` block under ``Practical Configuration Examples`` (exact heading): delete
    existing body paragraphs for that case, re-insert Purpose (non-bold), then one paragraph per
    assignment line, then a blank paragraph. Title paragraph is all-bold. Processes from last block
    upward so underlying ``python-docx`` references stay valid.
    """
    doc = Document(str(doc_path))
    blocks = collect_practical_configuration_example_blocks(doc)
    if not blocks:
        return 0
    for b in reversed(blocks):
        title = cast(Paragraph, b["title"])
        middle = cast(list[Paragraph], b["middle"])
        purposes = [p for p in middle if classify_practical_body_paragraph(p) == "purpose"]
        param_ps = [p for p in middle if classify_practical_body_paragraph(p) == "params"]
        others = [p for p in middle if classify_practical_body_paragraph(p) == "other"]
        other_texts = [(p.text or "").strip() for p in others if (p.text or "").strip()]
        param_lines: list[str] = []
        for p in param_ps:
            for ln in (p.text or "").splitlines():
                ts = ln.strip()
                if ts and is_param_assignment_line(ts):
                    param_lines.append(normalize_assignment_spacing(ts))
        purpose_text = "\n".join((p.text or "").strip() for p in purposes) if purposes else ""
        if middle:
            for p in sorted(middle, key=lambda q: paragraph_body_index(doc, q), reverse=True):
                delete_paragraph(p)
        anchor = title
        if purpose_text:
            anchor = insert_paragraph_after(anchor, purpose_text)
            set_paragraph_bold_all_runs(anchor, False)
        for txt in other_texts:
            anchor = insert_paragraph_after(anchor, txt)
            set_paragraph_bold_all_runs(anchor, False)
        for ln in param_lines:
            anchor = insert_paragraph_after(anchor, ln)
            set_paragraph_bold_all_runs(anchor, False)
        insert_paragraph_after(anchor, "")
        set_paragraph_bold_all_runs(title, True)
    doc.save(str(doc_path))
    return len(blocks)


def merge_dictionary_into_guidelines(doc_path: Path, mapping: dict[str, str]) -> tuple[int, int]:
    """
    For each parameter block whose name is in ``mapping``, yellow existing explanation runs and
    append dictionary text in new paragraph(s) after the last explanation paragraph (or after the
    title if there are none). Processes blocks from bottom to top so paragraph indices stay valid.
    Returns (blocks_merged, paragraphs_inserted).
    """
    doc = Document(str(doc_path))
    blocks = collect_guideline_param_blocks(doc)
    to_apply = [(ti, ei, p) for ti, ei, p in blocks if p in mapping and mapping[p].strip()]
    to_apply.sort(key=lambda x: x[0], reverse=True)
    merged = 0
    inserted = 0
    for title_i, expl_idxs, param in to_apply:
        paras = doc.paragraphs
        dict_text = mapping[param].strip()
        expl_ps = [paras[k] for k in expl_idxs if 0 <= k < len(paras)]
        highlight_explanation_paragraphs_yellow(expl_ps)
        anchor = expl_ps[-1] if expl_ps else paras[title_i]
        parts = [x.strip() for x in re.split(r"\n\s*\n", dict_text) if x.strip()]
        if not parts:
            parts = [dict_text]
        prev = anchor
        for part in parts:
            prev = insert_paragraph_after(prev, part)
            inserted += 1
        merged += 1
    doc.save(str(doc_path))
    return merged, inserted


def process(
    doc_path: Path,
    *,
    explicit_allowlist: frozenset[str] | None = None,
) -> tuple[int, int, int]:
    doc = Document(str(doc_path))
    if explicit_allowlist is not None:
        fields = set(explicit_allowlist)
    else:
        fields = reference_field_set(doc)
        if not fields:
            raise SystemExit("No Field column found in first table (use --explicit-allowlist).")

    paras = doc.paragraphs
    in_g = False
    i = 0
    blocks_touched = 0
    runs_removed = 0

    while i < len(paras):
        t = (paras[i].text or "").strip()
        low = t.lower()
        if low == "parameter configuration guidelines":
            in_g = True
            i += 1
            continue
        if in_g and low in STOP:
            break
        if not in_g:
            i += 1
            continue
        m = RE_PARAM_HEAD.match(t)
        if not m:
            i += 1
            continue

        param = m.group(1).upper()
        title_i = i
        j = i + 1
        expl_idxs: list[int] = []
        while j < len(paras):
            t2 = (paras[j].text or "").strip()
            l2 = t2.lower()
            if not t2:
                j += 1
                continue
            if l2 in STOP:
                break
            if RE_PARAM_HEAD.match(t2):
                break
            if RE_OPTIONS.match(t2):
                break
            expl_idxs.append(j)
            j += 1

        if param in fields:
            expl_ps = [paras[k] for k in expl_idxs]
            any_yellow = any(paragraph_has_yellow_run(p) for p in expl_ps)
            if any_yellow:
                title_p = paras[title_i]
                for p in expl_ps:
                    runs_removed += strip_yellow_runs(p)
                for p in reversed(expl_ps):
                    if not (p.text or "").strip():
                        delete_paragraph(p)
                highlight_title_yellow(title_p)
                blocks_touched += 1

        i = j

    blanks = remove_blank_paragraphs_in_guidelines(doc)
    doc.save(str(doc_path))
    return blocks_touched, runs_removed, blanks


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("docx", nargs="?", default=None, help="Path to DOCX (ignored if --directory is set)")
    ap.add_argument(
        "--directory",
        "-d",
        type=Path,
        default=None,
        help="Process all .docx in this folder in place.",
    )
    ap.add_argument(
        "--allowlist",
        choices=("table", "kernel", "ord-price-m"),
        default="table",
        help="Parameter name source for step 1 (default: first table Field column).",
    )
    ap.add_argument(
        "--explicit-allowlist",
        action="store_true",
        help="Deprecated: same as --allowlist kernel.",
    )
    ap.add_argument(
        "--mtime-today-before",
        metavar="HH:MM",
        default=None,
        help="With --directory: only files modified on today's local date with mtime before HH:MM.",
    )
    ap.add_argument(
        "--mtime-today-at",
        metavar="HH:MM",
        default=None,
        help="With --directory: only files modified during that clock minute today (local). Not with --mtime-today-before.",
    )
    ap.add_argument(
        "--merge-from-xlsx",
        type=Path,
        default=None,
        metavar="PATH",
        help="With --directory: merge dictionary explanations into guidelines (see module docstring).",
    )
    ap.add_argument(
        "--normalize-practical-examples",
        action="store_true",
        help="With --directory: normalize Practical Configuration Examples use-case blocks (see module docstring).",
    )
    args = ap.parse_args()
    if args.mtime_today_at and args.mtime_today_before:
        raise SystemExit("Use only one of --mtime-today-at and --mtime-today-before.")
    if args.normalize_practical_examples:
        if args.directory is None:
            raise SystemExit("--normalize-practical-examples requires --directory.")
        if (
            args.merge_from_xlsx is not None
            or args.mtime_today_at
            or args.mtime_today_before
            or args.explicit_allowlist
            or args.allowlist != "table"
        ):
            raise SystemExit(
                "Use only --normalize-practical-examples with --directory (no merge, mtime, or allowlist)."
            )
        d = args.directory
        if not d.is_dir():
            print("Not a directory:", d)
            return 1
        files = sorted(fp for fp in d.glob("*.docx") if not fp.name.startswith("~$"))
        if not files:
            print("No .docx files in:", d)
            return 0
        total = 0
        for path in files:
            try:
                n = normalize_practical_configuration_examples(path)
            except Exception as e:
                print(path.name, "ERROR", e)
                return 1
            total += n
            print(path.name, "use_cases_normalized", n)
        print("---")
        print("TOTAL files", len(files), "use_cases_blocks", total)
        return 0

    if args.merge_from_xlsx is not None:
        if args.directory is None:
            raise SystemExit("--merge-from-xlsx requires --directory.")
        if args.mtime_today_at or args.mtime_today_before:
            raise SystemExit("Do not combine --merge-from-xlsx with mtime filters.")
        d = args.directory
        if not d.is_dir():
            print("Not a directory:", d)
            return 1
        mapping = load_params_dictionary(args.merge_from_xlsx)
        print("Dictionary entries:", len(mapping), "from", args.merge_from_xlsx)
        files = sorted(fp for fp in d.glob("*.docx") if not fp.name.startswith("~$"))
        if not files:
            print("No .docx files in:", d)
            return 0
        total_m = total_ins = 0
        for path in files:
            try:
                m, ins = merge_dictionary_into_guidelines(path, mapping)
            except Exception as e:
                print(path.name, "ERROR", e)
                return 1
            total_m += m
            total_ins += ins
            print(path.name, "blocks_merged", m, "paras_inserted", ins)
        print("---")
        print("TOTAL files", len(files), "blocks_merged", total_m, "paras_inserted", total_ins)
        return 0

    mode = "kernel" if args.explicit_allowlist else args.allowlist
    allow: frozenset[str] | None = ALLOWLIST_BY_NAME[mode] if mode != "table" else None
    allow_msg = f"{mode} ({len(allow)} entries)" if allow is not None else "table[0] Field column"
    print("Allowlist:", allow_msg)

    if args.directory is not None:
        d = args.directory
        if not d.is_dir():
            print("Not a directory:", d)
            return 1
        files = sorted(fp for fp in d.glob("*.docx") if not fp.name.startswith("~$"))
        if args.mtime_today_at:
            matched = [fp for fp in files if mtime_today_at(fp, args.mtime_today_at)]
            skipped = len(files) - len(matched)
            print(
                "mtime filter: today at minute",
                args.mtime_today_at,
                "local -",
                len(matched),
                "of",
                len(files),
                "files",
                f"({skipped} skipped)" if skipped else "",
            )
            files = matched
        elif args.mtime_today_before:
            before = [fp for fp in files if mtime_today_before(fp, args.mtime_today_before)]
            skipped = len(files) - len(before)
            print(
                "mtime filter: today before",
                args.mtime_today_before,
                "local -",
                len(before),
                "of",
                len(files),
                "files",
                f"({skipped} skipped)" if skipped else "",
            )
            files = before
        if not files:
            print("No .docx files to process in:", d)
            return 0
        sum_bt = sum_rr = sum_bl = 0
        for path in files:
            bt, rr, blanks = process(path, explicit_allowlist=allow)
            sum_bt += bt
            sum_rr += rr
            sum_bl += blanks
            print(path.name, "blocks_updated", bt, "yellow_runs_removed", rr, "blank_paras_removed", blanks)
        print("---")
        print("TOTAL files", len(files), "blocks", sum_bt, "runs_removed", sum_rr, "blanks", sum_bl)
        return 0

    path = Path(args.docx) if args.docx else DEFAULT_DOC
    if not path.is_file():
        print("Missing:", path)
        return 1

    bt, rr, blanks = process(path, explicit_allowlist=allow)
    print(path.name, "blocks_updated", bt, "yellow_runs_removed", rr, "blank_paras_removed", blanks)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
