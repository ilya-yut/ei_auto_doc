"""
Normalize 'Practical Configuration Examples' in Parts1_4_enriched DOCX (Parts 2–4 when run as __main__; adjust PARTS below to include Part 1 if needed).

Rules (user corrections):
- Canonical section title: Practical Configuration Examples (bold on title paragraph).
- Each use case: Use Case N: <title> (bold) / Purpose: <text> (not bold) / assignment lines (not bold).
- One paragraph per line; blank paragraph after each use case (except no trailing blank before next section).
- Equality: spaces around = ; comparisons NAME OP VALUE with spaces around OP.
- Renumber Use Case 1..N. No extra intro text.
- Light blue paragraph shading ONLY on paragraphs whose normalized content changed vs pre-edit snapshot
  (or new lines / blank separators not present at same index).

In-place save. Stop section at Heading 1/2 or known section titles.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(r"c:\vibe code dev\ei_auto_doc")
ENRICHED = ROOT / "ei docs for analysis" / "EI docs" / "temp" / "Parts1_4_enriched"

CANONICAL_TITLE = "Practical Configuration Examples"
LIGHT_BLUE = "D9E9F7"

STOP_PARA_TEXT = frozenset(
    {
        "ei function structure",
        "abap code",
        "default values",
        "parameter relationship",
        "parameter relationships",
        "parameter configuration guidelines",
    }
)

RE_PRACTICAL_TITLE = re.compile(
    r"^\s*Practical\s+(Example\s+of\s+Parameter\s+Configuration|Configuration\s+Examples)\s*$",
    re.I,
)
RE_USE_CASE_START = re.compile(r"^Use Case\s+\d+\s*:\s*", re.I)


def _remove_shade(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:shd"):
            p_pr.remove(child)


def _shade_paragraph(paragraph: Paragraph, fill_hex: str = LIGHT_BLUE) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:shd"):
            p_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def _delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_paragraph_after(
    paragraph: Paragraph, text: str, style_name: str | None = None, bold: bool = False
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass
    r = new_para.add_run(text)
    r.bold = bool(bold)
    return new_para


def _is_stop_paragraph(p: Paragraph) -> bool:
    t = (p.text or "").strip()
    if not t:
        return False
    st = (p.style and p.style.name) or ""
    if st == "Heading 2":
        return True
    if st == "Heading 1":
        return True
    low = t.lower()
    if low in STOP_PARA_TEXT:
        return True
    return False


def _find_practical_section(doc: Document) -> tuple[int, int] | None:
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        if not (p.text or "").strip():
            continue
        if RE_PRACTICAL_TITLE.match((p.text or "").strip()):
            start = i
            break
    if start is None:
        return None
    j = start + 1
    while j < len(paras):
        if _is_stop_paragraph(paras[j]):
            break
        j += 1
    return start, j


def _normalize_setting_line(line: str) -> str | None:
    s = line.strip()
    m_eq = re.match(r"^([A-Z][A-Z0-9_]*)\s*=\s*(.*)$", s)
    if m_eq:
        return f"{m_eq.group(1)} = {m_eq.group(2).strip()}"
    m_cmp = re.match(r"^([A-Z][A-Z0-9_]*)\s*(>=|<=|<>|>|<)\s*(.*)$", s)
    if m_cmp:
        name, op, val = m_cmp.group(1), m_cmp.group(2), m_cmp.group(3).strip()
        return f"{name} {op} {val}"
    return None


def _norm_compare_line(s: str) -> str:
    t = s.strip()
    if not t:
        return ""
    low = t.lower()
    if low.startswith("purpose:"):
        return "Purpose:" + t.split(":", 1)[1].strip()
    sl = _normalize_setting_line(t)
    if sl:
        return sl
    return t


def _flatten_old_body_lines(doc: Document, start: int, end: int) -> list[str]:
    """One entry per logical line (splitting paragraph newlines), empty paras as ''."""
    out: list[str] = []
    for k in range(start + 1, end):
        raw = doc.paragraphs[k].text or ""
        if not raw.strip():
            out.append("")
            continue
        parts = raw.split("\n")
        for p in parts:
            out.append(p.strip())
    return out


def _parse_use_cases(full_text: str) -> list[dict[str, str | list[str]]] | None:
    full_text = full_text.strip()
    if not full_text:
        return []

    starts = [m.start() for m in re.finditer(r"(?m)^Use Case\s+\d+\s*:\s*", full_text)]
    if not starts:
        return None

    blocks: list[str] = []
    for idx, st in enumerate(starts):
        ed = starts[idx + 1] if idx + 1 < len(starts) else len(full_text)
        blocks.append(full_text[st:ed].strip())

    out: list[dict[str, str | list[str]]] = []
    for block in blocks:
        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue
        m0 = re.match(r"(?i)^Use Case\s+\d+\s*:\s*(.*)$", lines[0])
        if not m0:
            return None
        title = m0.group(1).strip()
        purpose_parts: list[str] = []
        params: list[str] = []
        i = 1
        while i < len(lines):
            ln = lines[i]
            low = ln.lower()
            if low.startswith("purpose:"):
                purpose_parts.append(ln.split(":", 1)[1].strip())
                i += 1
                while i < len(lines):
                    if _normalize_setting_line(lines[i]) or lines[i].lower().startswith("purpose:"):
                        break
                    purpose_parts.append(lines[i])
                    i += 1
                continue
            sl = _normalize_setting_line(ln)
            if sl:
                params.append(sl)
                i += 1
                continue
            if params:
                return None
            purpose_parts.append(ln)
            i += 1

        purpose = " ".join(purpose_parts).strip() or title
        out.append({"title": title, "purpose": purpose, "params": params})
    return out


def _section_body_text(doc: Document, start: int, end: int) -> str:
    parts: list[str] = []
    for k in range(start + 1, end):
        t = (doc.paragraphs[k].text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _build_section_lines(use_cases: list[dict[str, str | list[str]]]) -> list[str]:
    lines: list[str] = []
    n = len(use_cases)
    for i, uc in enumerate(use_cases, start=1):
        lines.append(f"Use Case {i}: {uc['title']}")
        lines.append(f"Purpose: {uc['purpose']}")
        for pl in uc["params"]:
            lines.append(pl)
        if i < n:
            lines.append("")
    return lines


def _line_needs_shade(new_line: str, idx: int, old_lines: list[str]) -> bool:
    if idx >= len(old_lines):
        return True
    return _norm_compare_line(new_line) != _norm_compare_line(old_lines[idx])


def _runs_all_bold(paragraph: Paragraph) -> bool:
    runs = paragraph.runs
    if not runs:
        return False
    return all(r.bold is True for r in runs)


def _clear_section_shading(doc: Document, start: int, end: int) -> bool:
    """Remove w:shd from title + body paragraphs in [start, end). Return True if anything removed."""
    changed = False
    for k in range(start, end):
        p = doc.paragraphs[k]
        p_pr = p._p.pPr
        if p_pr is None:
            continue
        if any(c.tag == qn("w:shd") for c in p_pr):
            _remove_shade(p)
            changed = True
    return changed


def _ensure_bold_use_cases_only(doc: Document, start: int, end: int) -> bool:
    """Set bold on section title and Use Case title lines only; return True if any run changed."""
    changed = False
    tp = doc.paragraphs[start]
    if (tp.text or "").strip() == CANONICAL_TITLE and not _runs_all_bold(tp):
        txt = (tp.text or "").strip()
        tp.clear()
        r = tp.add_run(txt)
        r.bold = True
        changed = True
    for k in range(start + 1, end):
        p = doc.paragraphs[k]
        t = (p.text or "").strip()
        if not t:
            continue
        if RE_USE_CASE_START.match(t) and not _runs_all_bold(p):
            p.clear()
            r = p.add_run(t)
            r.bold = True
            changed = True
    return changed


def process_document(path: Path) -> tuple[bool, str]:
    doc = Document(str(path))
    loc = _find_practical_section(doc)
    if loc is None:
        return False, "no practical section"
    start, end = loc
    title_para = doc.paragraphs[start]
    old_title = (title_para.text or "").strip()
    old_body_lines = _flatten_old_body_lines(doc, start, end)
    body_text = _section_body_text(doc, start, end)

    parsed = _parse_use_cases(body_text)
    if parsed is None:
        return False, "unparseable use cases"

    new_lines = _build_section_lines(parsed)

    title_changed = old_title.strip() != CANONICAL_TITLE

    same_body = len(new_lines) == len(old_body_lines) and all(
        _norm_compare_line(a) == _norm_compare_line(b) for a, b in zip(new_lines, old_body_lines)
    )
    if not title_changed and same_body:
        if _ensure_bold_use_cases_only(doc, start, end):
            doc.save(str(path))
            return True, "bold-only"
        if _clear_section_shading(doc, start, end):
            doc.save(str(path))
            return True, "cleared-residual-shading"
        return False, "already normalized"

    body_style = None
    if end > start + 1:
        try:
            body_style = doc.paragraphs[start + 1].style.name
        except Exception:
            body_style = None
    if not body_style:
        body_style = "Normal"

    _remove_shade(title_para)
    for k in range(start + 1, end):
        _remove_shade(doc.paragraphs[k])

    for idx in range(end - 1, start, -1):
        _delete_paragraph(doc.paragraphs[idx])

    title_para = doc.paragraphs[start]
    title_para.clear()
    tr = title_para.add_run(CANONICAL_TITLE)
    tr.bold = True
    if title_changed:
        _shade_paragraph(title_para)

    anchor = title_para
    for i, ln in enumerate(new_lines):
        bold = bool(ln.strip() and RE_USE_CASE_START.match(ln.strip()))
        anchor = insert_paragraph_after(anchor, ln, style_name=body_style, bold=bold)
        if _line_needs_shade(ln, i, old_body_lines):
            _shade_paragraph(anchor)

    doc.save(str(path))
    return True, "updated"


def main() -> None:
    parts = ["Part 2", "Part 3", "Part 4"]
    total_changed = 0
    total_skipped = 0
    errors: list[str] = []
    for part in parts:
        d = ENRICHED / part
        if not d.is_dir():
            print(f"Missing: {d}")
            continue
        changed = 0
        skipped = 0
        for fp in sorted(d.glob("*.docx")):
            if fp.name.startswith("~$"):
                continue
            try:
                ok, msg = process_document(fp)
                if ok:
                    changed += 1
                    print(f"OK {part}/{fp.name}: {msg}")
                else:
                    skipped += 1
                    if msg not in ("no practical section", "already normalized"):
                        errors.append(f"{part}/{fp.name}: {msg}")
            except Exception as e:
                errors.append(f"{part}/{fp.name}: {e!r}")
        print(f"{part} — updated: {changed}, skipped/unchanged: {skipped}")
        total_changed += changed
        total_skipped += skipped
    print(f"All parts — updated: {total_changed}, skipped/unchanged: {total_skipped}")
    for e in errors[:80]:
        print("WARN", e)


if __name__ == "__main__":
    main()
