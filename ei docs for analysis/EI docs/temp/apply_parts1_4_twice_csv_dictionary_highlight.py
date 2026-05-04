"""
For each DOCX under temp/Parts1_4_enriched/Part 1..Part 4 (in place, no copies):

- Use params_dictionary_from_twice_csv.xlsx (sheet dictionary).
- In "Parameter Configuration Guidelines", for parameters that exist in the dictionary:
  - Highlight existing explanation paragraphs in yellow
  - Insert the dictionary explanation below the highlighted block

Does not create folders or new DOCX paths; only saves files that were modified.
Re-runs skip parameters whose new dictionary text already appears (no duplicate stacks).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook


ROOT = Path(r"c:\vibe code dev\ei_auto_doc")
TEMP_DIR = ROOT / "ei docs for analysis" / "EI docs" / "temp"
ENRICHED_BASE = TEMP_DIR / "Parts1_4_enriched"

DICT_XLSX = TEMP_DIR / "params_dictionary_from_twice_csv.xlsx"

RE_PARAM_HEAD = re.compile(r"^([A-Z][A-Z0-9_]*)\s*\(.*\)\s*:?\s*$")
RE_OPTIONS = re.compile(r"^(?:\*?\*?)?[A-Z][A-Z0-9_]*\s*Options:\s*(?:\*?\*?)?$", re.I)
STOP_TITLES = {
    "parameter relationship",
    "parameter relationships",
    "default values",
    "practical example of parameter configuration",
    "practical configuration examples",
    "ei function structure",
    "abap code",
}


def insert_paragraph_after(paragraph: Paragraph, text: str, style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def _header_col_map(ws) -> dict[str, int]:
    out: dict[str, int] = {}
    mc = ws.max_column or 0
    for c in range(1, mc + 1):
        v = ws.cell(1, c).value
        if v:
            out[str(v).strip().lower()] = c
    return out


def load_dict_map(path: Path) -> dict[str, str]:
    wb = load_workbook(path, read_only=True)
    ws = wb["dictionary"] if "dictionary" in wb.sheetnames else wb[wb.sheetnames[0]]
    hm = _header_col_map(ws)
    expl_col = hm.get("sap canonical explanation")
    if expl_col is None:
        expl_col = hm.get("suggested/corrected explanation")
    if expl_col is None and (ws.max_column or 0) >= 3:
        expl_col = 3
    if expl_col is None:
        expl_col = 2

    out: dict[str, str] = {}
    for r in range(2, ws.max_row + 1):
        p = ws.cell(r, 1).value
        if not p:
            continue
        param = str(p).strip().upper()
        raw = ws.cell(r, expl_col).value
        expl = str(raw).strip() if raw else ""
        if param and expl:
            out[param] = expl
    wb.close()
    return out


def split_paragraphs(text: str) -> list[str]:
    parts = [x.strip() for x in text.split("\n\n") if x.strip()]
    return parts if parts else [text.strip()]


def process_docx(doc_path: Path, explain_map: dict[str, str]) -> tuple[bool, int]:
    doc = Document(str(doc_path))
    paras = doc.paragraphs
    changed = False
    inserts = 0

    in_guidelines = False
    i = 0
    while i < len(paras):
        p = paras[i]
        txt = (p.text or "").strip()
        low = txt.lower()

        if low == "parameter configuration guidelines":
            in_guidelines = True
            i += 1
            continue

        if in_guidelines and low in STOP_TITLES:
            in_guidelines = False
            i += 1
            continue

        if not in_guidelines:
            i += 1
            continue

        m = RE_PARAM_HEAD.match(txt)
        if not m:
            i += 1
            continue

        param = m.group(1).upper()
        new_expl = explain_map.get(param)
        if not new_expl:
            i += 1
            continue

        j = i + 1
        explanation_idxs: list[int] = []
        found_same_new_text = False
        pieces_check = split_paragraphs(new_expl)
        while j < len(paras):
            pj = paras[j]
            t = (pj.text or "").strip()
            l = t.lower()
            if not t:
                j += 1
                continue
            if l in STOP_TITLES:
                break
            if RE_PARAM_HEAD.match(t):
                break
            if RE_OPTIONS.match(t):
                break
            if t == new_expl.strip() or t in pieces_check:
                found_same_new_text = True
            explanation_idxs.append(j)
            j += 1

        if found_same_new_text:
            i = j
            continue

        pieces = split_paragraphs(new_expl)
        if explanation_idxs:
            for idx in explanation_idxs:
                para = paras[idx]
                if not para.runs:
                    para.add_run("")
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            anchor = paras[explanation_idxs[-1]]
            style_name = anchor.style.name if anchor.style else None
            cur = anchor
            for piece in pieces:
                cur = insert_paragraph_after(cur, piece, style_name=style_name)
            changed = True
            inserts += len(pieces)
            paras = doc.paragraphs
            i = j + 1
        else:
            anchor = paras[i]
            style_name = anchor.style.name if anchor.style else None
            cur = anchor
            for piece in pieces:
                cur = insert_paragraph_after(cur, piece, style_name=style_name)
            changed = True
            inserts += len(pieces)
            paras = doc.paragraphs
            i += 2

    if changed:
        doc.save(str(doc_path))
    return changed, inserts


def main() -> None:
    explain_map = load_dict_map(DICT_XLSX)

    parts = ["Part 1", "Part 2", "Part 3", "Part 4"]
    total_files = 0
    changed_files = 0
    total_inserts = 0

    for part in parts:
        work_dir = ENRICHED_BASE / part
        if not work_dir.exists():
            print(f"Skip missing folder: {work_dir}")
            continue
        for fp in sorted(work_dir.glob("*.docx")):
            if fp.name.startswith("~$"):
                continue
            total_files += 1
            changed, ins = process_docx(fp, explain_map)
            if changed:
                changed_files += 1
                total_inserts += ins

    print(f"Dictionary: {DICT_XLSX}")
    print(f"In-place root: {ENRICHED_BASE}")
    print(f"Dictionary parameters with explanation: {len(explain_map)}")
    print(f"DOCX scanned (Part 1-4): {total_files}")
    print(f"Files modified (had matching params): {changed_files}")
    print(f"Paragraph inserts: {total_inserts}")


if __name__ == "__main__":
    main()
